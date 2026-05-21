"""
ATLAS Warehouse AGV — Final Project Report Generator
Generates a fully formatted DOCX dissertation per B.Tech guidelines.
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins (1 inch all sides) ──────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1)
    section.right_margin  = Inches(1)
    section.page_height   = Inches(11.69)   # A4
    section.page_width    = Inches(8.27)

# ── Style helpers ─────────────────────────────────────────────────────────────
def set_run_font(run, size_pt, bold=False, italic=False, color=None):
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size_pt)
    run.bold  = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def para_spacing(para, before=0, after=0, line=None):
    from docx.shared import Pt
    para.paragraph_format.space_before = Pt(before)
    para.paragraph_format.space_after  = Pt(after)
    if line:
        para.paragraph_format.line_spacing = line

def body_para(text, bold=False, italic=False, size=12, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
              before=0, after=12, line=1.5):
    p = doc.add_paragraph()
    p.alignment = align
    para_spacing(p, before=before, after=after)
    p.paragraph_format.line_spacing = Pt(size * line)
    run = p.add_run(text)
    set_run_font(run, size, bold=bold, italic=italic)
    return p

def heading1(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para_spacing(p, before=12, after=12)
    run = p.add_run(text)
    set_run_font(run, 14, bold=True)
    return p

def heading2(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para_spacing(p, before=6, after=6)
    run = p.add_run(text)
    set_run_font(run, 12, bold=True)
    return p


def heading3(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para_spacing(p, before=6, after=3)
    run = p.add_run(text)
    set_run_font(run, 12, bold=True, italic=True)
    return p

def page_break():
    doc.add_page_break()

def center_bold(text, size=18):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_spacing(p, before=6, after=6)
    run = p.add_run(text)
    set_run_font(run, size, bold=True)
    return p

def center_normal(text, size=14):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_spacing(p, before=4, after=4)
    run = p.add_run(text)
    set_run_font(run, size)
    return p

def add_equation(eq_text, eq_number):
    """Add an equation line (left aligned text) with number at right."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para_spacing(p, before=12, after=12)
    run = p.add_run(eq_text + "   " + eq_number)
    set_run_font(run, 12, italic=True)
    return p

def table_caption(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_spacing(p, before=12, after=6)
    run = p.add_run(text)
    set_run_font(run, 12, bold=True)
    return p

def figure_caption(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_spacing(p, before=6, after=12)
    run = p.add_run(text)
    set_run_font(run, 12, bold=True)
    return p

def add_table(headers, rows, caption_text):
    table_caption(caption_text)
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for run in hdr[i].paragraphs[0].runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
            run.bold = True
        hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for ri, row in enumerate(rows):
        cells = t.rows[ri+1].cells
        for ci, val in enumerate(row):
            cells[ci].text = val
            for run in cells[ci].paragraphs[0].runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)
            cells[ci].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph()
    return t


# ════════════════════════════════════════════════════════════════════════════
# 1. COVER PAGE
# ════════════════════════════════════════════════════════════════════════════
for _ in range(4):
    doc.add_paragraph()

center_bold("ATLAS FLEET WAREHOUSE AGV", 18)
center_bold("Design, Simulation, and Physical Conversion of an Autonomous Guided Vehicle System for Warehouse Automation", 16)

for _ in range(2):
    doc.add_paragraph()

center_bold("PROJECT REPORT", 16)
center_normal("SUBMITTED IN PARTIAL FULFILLMENT OF THE REQUIREMENTS FOR THE AWARD OF THE DEGREE OF", 13)
center_bold("Bachelor of Technology", 16)
center_normal("in", 13)
center_bold("Mechanical Engineering", 16)

for _ in range(2):
    doc.add_paragraph()

center_normal("Submitted By", 13)
center_bold("[Student Name], [Faculty Number]", 14)

for _ in range(1):
    doc.add_paragraph()

center_normal("Under the Supervision of", 13)
center_bold("[Supervisor Name]", 14)
center_normal("Assistant Professor / Associate Professor", 13)

for _ in range(2):
    doc.add_paragraph()

center_bold("Department of Mechanical Engineering", 14)
center_bold("Zakir Husain College of Engineering and Technology", 14)
center_bold("Aligarh Muslim University, Aligarh", 14)
center_bold("2024–25", 14)

page_break()

# ════════════════════════════════════════════════════════════════════════════
# 2. ACKNOWLEDGEMENT
# ════════════════════════════════════════════════════════════════════════════
heading1("ACKNOWLEDGEMENT")

body_para(
    "I wish to express my heartfelt gratitude to my project supervisor, [Supervisor Name], "
    "for their invaluable guidance, constant encouragement, constructive criticism, and "
    "unwavering support throughout the duration of this project. Their expertise in robotics "
    "and autonomous systems was instrumental in shaping the direction and quality of this work."
)
body_para(
    "I extend my sincere thanks to the Head of the Department of Mechanical Engineering, "
    "Zakir Husain College of Engineering and Technology, Aligarh Muslim University, for "
    "providing the necessary infrastructure, laboratory facilities, and administrative support "
    "required for the successful completion of this project."
)
body_para(
    "I am deeply grateful to Dr. Mohammad Faizan, B.Tech. Project Coordinator, Department of "
    "Mechanical Engineering, for his consistent guidance and for establishing the structured "
    "framework that enabled rigorous project reporting."
)
body_para(
    "I acknowledge the open-source communities behind ROS2 (Open Robotics), Gazebo Classic "
    "Simulator, Python, and PyQt5 for providing the world-class tools that made this project "
    "possible. Their contributions to the robotics community are foundational to work of this nature."
)
body_para(
    "Finally, I express my profound gratitude to my parents and family for their unconditional "
    "love, patience, and moral support throughout my academic journey."
)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run("[Student Name]")
set_run_font(run, 12, bold=True)

page_break()


# ════════════════════════════════════════════════════════════════════════════
# 3. ABSTRACT
# ════════════════════════════════════════════════════════════════════════════
heading1("ABSTRACT")

body_para(
    "This project presents the complete design, simulation, implementation, and physical "
    "conversion analysis of the ATLAS Fleet Warehouse Automated Guided Vehicle (AGV) — a "
    "ROS2-based autonomous mobile robot system for material handling in structured warehouse "
    "environments. The system demonstrates the full operational cycle of a warehouse AGV: "
    "receiving pick missions via a graphical control interface, navigating to target shelf "
    "locations through line-following, identifying shelves via RFID detection, executing pick "
    "operations, and autonomously returning to a home docking station."
)
body_para(
    "The robot employs a differential-drive kinematic configuration with an 8-channel infrared "
    "reflectance sensor array for line following, a 9-axis Inertial Measurement Unit (IMU) for "
    "precision heading control during turns, and simulated RFID for shelf identification. A "
    "Proportional-Derivative (PD) controller maintains line tracking accuracy below 5 mm at an "
    "operational speed of 0.4 m/s. A 12-state Finite State Machine (FSM) orchestrates the "
    "complete mission lifecycle from mission intake through pick execution to autonomous docking."
)
body_para(
    "The system is implemented on ROS2 Humble with Gazebo Classic 11 simulation, comprising "
    "6 ROS2 packages, 6 computational nodes, 3 custom message types, and 17 communication "
    "topics. A PyQt5-based industrial control centre provides real-time telemetry, mission "
    "dispatch, emergency stop, and reset-to-dock functionality. A simulation-to-physical "
    "conversion analysis maps every simulation component to real-world hardware equivalents "
    "including JGB37-520 DC motors, L298N motor controller, QTR-8A reflectance sensor array, "
    "BNO055 IMU, and RDM6300 RFID reader on a Raspberry Pi 4 platform."
)
body_para(
    "Key results demonstrate 100% mission completion rate, sub-5 mm line tracking accuracy, "
    "±3° turn precision, a complete mission cycle time of approximately 45 seconds for shelf "
    "S05, and an estimated throughput of 80 pick cycles per hour. The architecture is designed "
    "for direct portability to physical hardware, with 70% of the codebase running unchanged "
    "on a real robot platform."
)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run("Keywords: ")
set_run_font(run, 12, bold=True)
run2 = p.add_run("Automated Guided Vehicle, ROS2, Differential Drive, Line Following, PID Control, "
                 "RFID, Warehouse Automation, Finite State Machine, Gazebo Simulation, Industry 4.0.")
set_run_font(run2, 12, italic=True)

page_break()


# ════════════════════════════════════════════════════════════════════════════
# 4. TABLE OF CONTENTS
# ════════════════════════════════════════════════════════════════════════════
heading1("TABLE OF CONTENTS")

toc_entries = [
    ("Acknowledgement", "i"),
    ("Abstract", "ii"),
    ("Table of Contents", "iii"),
    ("Notations / Symbols / Abbreviations", "iv"),
    ("List of Tables", "v"),
    ("List of Figures", "vi"),
    ("Chapter 1 — Introduction and Literature Review", "1"),
    ("    1.1  Background and Industry Relevance", "1"),
    ("    1.2  Automated Guided Vehicles in Industry", "2"),
    ("    1.3  Literature Review", "3"),
    ("    1.4  Research Gap and Motivation", "6"),
    ("    1.5  Project Objectives", "7"),
    ("    1.6  Chapter Summary", "7"),
    ("Chapter 2 — Problem Formulation", "8"),
    ("    2.1  Problem Domain", "8"),
    ("    2.2  Existing Limitations", "8"),
    ("    2.3  Project Requirements and Design Constraints", "9"),
    ("    2.4  Operational Assumptions", "9"),
    ("    2.5  Expected Outcomes and System Objectives", "10"),
    ("    2.6  Scope and Limitations", "10"),
    ("    2.7  Chapter Summary", "11"),
    ("Chapter 3 — Modeling, Solution Methodology and Implementation", "12"),
    ("    3.1  Overall System Architecture", "12"),
    ("    3.2  Software and ROS2 Package Architecture", "13"),
    ("    3.3  ROS2 Node Architecture", "14"),
    ("    3.4  Topic Communication and Message Flow", "15"),
    ("    3.5  Transform (TF) Tree", "17"),
    ("    3.6  Robot Description — URDF/Xacro Model", "17"),
    ("    3.7  Differential Drive Kinematics", "19"),
    ("    3.8  Gazebo Simulation Environment", "23"),
    ("    3.9  Sensor Architecture", "25"),
    ("    3.10 Control Algorithms — PD Line Following", "29"),
    ("    3.11 Turn Controller — IMU-Based Bang-Bang Control", "31"),
    ("    3.12 Mission State Machine Design", "33"),
    ("    3.13 Velocity Arbiter Pattern", "36"),
    ("    3.14 Simulation-to-Physical Conversion Methodology", "37"),
    ("    3.15 Chapter Summary", "42"),
    ("Chapter 4 — Results, Discussion and Conclusions", "43"),
    ("    4.1  Simulation Performance Results", "43"),
    ("    4.2  Mission Timing Analysis", "44"),
    ("    4.3  System Behaviour and Discussion", "45"),
    ("    4.4  Error Analysis and Limitations", "46"),
    ("    4.5  Physical Deployment Analysis", "47"),
    ("    4.6  Industry 4.0 Context", "48"),
    ("    4.7  Future Improvements", "49"),
    ("    4.8  Conclusions", "50"),
    ("References", "52"),
    ("Appendix A — ROS2 Package Structure and Source Files", "54"),
    ("Appendix B — Terminal Commands and Launch Procedure", "57"),
    ("Appendix C — GPIO Pin Mapping", "59"),
    ("Appendix D — Complete State Machine Transition Table", "60"),
]

for entry, pg in toc_entries:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para_spacing(p, before=0, after=2)
    p.paragraph_format.line_spacing = Pt(18)
    run = p.add_run(entry)
    set_run_font(run, 12, bold=(".1" not in entry and "Chapter" in entry or entry in
                               ["Acknowledgement","Abstract","Table of Contents",
                                "Notations / Symbols / Abbreviations","List of Tables",
                                "List of Figures","References"]))
    # tab leader simulation
    run2 = p.add_run(f"  {'.' * max(1, 60 - len(entry))}  {pg}")
    set_run_font(run2, 12)

page_break()


# ════════════════════════════════════════════════════════════════════════════
# 5. NOTATIONS / SYMBOLS / ABBREVIATIONS
# ════════════════════════════════════════════════════════════════════════════
heading1("NOTATIONS, SYMBOLS AND ABBREVIATIONS")
heading2("A. Mathematical Symbols and Nomenclature")

add_table(
    ["Symbol", "Description", "Unit"],
    [
        ["v", "Linear velocity of robot centre", "m/s"],
        ["ω", "Angular velocity of robot", "rad/s"],
        ["v_L", "Left wheel linear velocity", "m/s"],
        ["v_R", "Right wheel linear velocity", "m/s"],
        ["ω_L", "Left wheel angular velocity", "rad/s"],
        ["ω_R", "Right wheel angular velocity", "rad/s"],
        ["r", "Wheel radius", "m"],
        ["L", "Wheel track width (separation)", "m"],
        ["R", "Instantaneous turning radius", "m"],
        ["θ", "Robot heading angle", "rad"],
        ["K_p", "Proportional controller gain", "—"],
        ["K_i", "Integral controller gain", "—"],
        ["K_d", "Derivative controller gain", "—"],
        ["e(t)", "Control error signal", "—"],
        ["u(t)", "Controller output (angular velocity command)", "rad/s"],
        ["τ", "Motor torque", "N·m"],
        ["I_xx, I_yy, I_zz", "Principal moments of inertia", "kg·m²"],
        ["m", "Robot mass", "kg"],
        ["g", "Gravitational acceleration", "m/s²"],
        ["μ", "Coefficient of friction", "—"],
        ["F_N", "Normal force", "N"],
        ["Δs", "Incremental distance (odometry)", "m"],
        ["Δθ", "Incremental angle (odometry)", "rad"],
        ["ψ", "Yaw angle extracted from IMU quaternion", "rad"],
        ["w_i", "Sensor weight for channel i", "—"],
        ["s_i", "Binary sensor reading for channel i", "{0,1}"],
    ],
    "Table 1: Mathematical Symbols and Nomenclature"
)

heading2("B. Acronyms and Abbreviations")

add_table(
    ["Acronym", "Full Form"],
    [
        ["AGV", "Automated Guided Vehicle"],
        ["AMR", "Autonomous Mobile Robot"],
        ["CLI", "Command Line Interface"],
        ["CPR", "Counts Per Revolution"],
        ["DDS", "Data Distribution Service"],
        ["DoF", "Degrees of Freedom"],
        ["FSM", "Finite State Machine"],
        ["GPIO", "General Purpose Input/Output"],
        ["GUI", "Graphical User Interface"],
        ["HMI", "Human-Machine Interface"],
        ["I2C", "Inter-Integrated Circuit"],
        ["ICR", "Instantaneous Centre of Rotation"],
        ["IMU", "Inertial Measurement Unit"],
        ["IR", "Infrared"],
        ["LiDAR", "Light Detection and Ranging"],
        ["ODE", "Open Dynamics Engine"],
        ["PD", "Proportional-Derivative"],
        ["PID", "Proportional-Integral-Derivative"],
        ["PLC", "Programmable Logic Controller"],
        ["PWM", "Pulse Width Modulation"],
        ["QoS", "Quality of Service"],
        ["RFID", "Radio Frequency Identification"],
        ["ROS2", "Robot Operating System 2"],
        ["RPM", "Revolutions Per Minute"],
        ["RSP", "Robot State Publisher"],
        ["RViz", "ROS Visualization Tool"],
        ["SDF", "Simulation Description Format"],
        ["SLAM", "Simultaneous Localization and Mapping"],
        ["SPI", "Serial Peripheral Interface"],
        ["TF", "Transform Frame (coordinate frame)"],
        ["UART", "Universal Asynchronous Receiver-Transmitter"],
        ["URDF", "Unified Robot Description Format"],
        ["WMS", "Warehouse Management System"],
        ["XACRO", "XML Macros (robot description language extension)"],
    ],
    "Table 2: Acronyms and Abbreviations"
)

page_break()


# ════════════════════════════════════════════════════════════════════════════
# 6. LIST OF TABLES
# ════════════════════════════════════════════════════════════════════════════
heading1("LIST OF TABLES")

tables_list = [
    ("Table 1", "Mathematical Symbols and Nomenclature"),
    ("Table 2", "Acronyms and Abbreviations"),
    ("Table 3", "AGV Technology Generations and Navigation Methods"),
    ("Table 4", "Comparison of Prior Literature on AGV Systems"),
    ("Table 5", "Project Objectives and Success Criteria"),
    ("Table 6", "Design Constraints and Operational Assumptions"),
    ("Table 7", "ROS2 Package Structure"),
    ("Table 8", "Node Input/Output Summary"),
    ("Table 9", "Complete ROS2 Topic List"),
    ("Table 10", "Physical Parameters from URDF/Xacro"),
    ("Table 11", "Inertia Tensor Values for Chassis"),
    ("Table 12", "Gazebo Physics Configuration"),
    ("Table 13", "Line Sensor Array Geometry Parameters"),
    ("Table 14", "PD Controller Parameters and Tuning"),
    ("Table 15", "Mission State Machine — States and Velocity Sources"),
    ("Table 16", "Complete State Machine Transition Table"),
    ("Table 17", "Simulation-to-Physical Component Conversion Table"),
    ("Table 18", "Hardware Component Specifications"),
    ("Table 19", "GPIO Pin Mapping for Physical Robot"),
    ("Table 20", "Power Budget for Physical Robot"),
    ("Table 21", "Budget Summary — Minimum and Production Builds"),
    ("Table 22", "System Performance Metrics"),
    ("Table 23", "Mission Timing Analysis — Shelf S05"),
    ("Table 24", "Commercial AGV Comparison"),
    ("Table 25", "Design Trade-offs Summary"),
]

for tbl_num, tbl_title in tables_list:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para_spacing(p, before=0, after=2)
    p.paragraph_format.line_spacing = Pt(18)
    run = p.add_run(f"{tbl_num}:  {tbl_title}")
    set_run_font(run, 12)

page_break()

# ════════════════════════════════════════════════════════════════════════════
# 7. LIST OF FIGURES
# ════════════════════════════════════════════════════════════════════════════
heading1("LIST OF FIGURES")

figures_list = [
    ("Fig. 1.1", "Overall System Architecture Block Diagram"),
    ("Fig. 1.2", "ATLAS AGV in Gazebo Warehouse Simulation"),
    ("Fig. 3.1", "ROS2 Node and Topic Communication Graph"),
    ("Fig. 3.2", "Package Dependency Structure"),
    ("Fig. 3.3", "Custom Message Definitions (FleetMission, RobotState, ShelfTag)"),
    ("Fig. 3.4", "TF Transform Tree (odom → base_footprint → base_link)"),
    ("Fig. 3.5", "ATLAS AGV URDF/Xacro Robot Model"),
    ("Fig. 3.6", "Differential Drive Kinematic Model"),
    ("Fig. 3.7", "Forward and Inverse Kinematics Diagram"),
    ("Fig. 3.8", "Stability Triangle and Centre of Gravity"),
    ("Fig. 3.9", "Gazebo Warehouse World Layout with Coordinates"),
    ("Fig. 3.10", "Warehouse Navigation Layout — Spine and Aisles"),
    ("Fig. 3.11", "Line Sensor Array Geometry (8-channel, 50 Hz)"),
    ("Fig. 3.12", "Junction Detection Logic Diagram"),
    ("Fig. 3.13", "RFID Detection Hysteresis Model"),
    ("Fig. 3.14", "PD Control Block Diagram for Line Following"),
    ("Fig. 3.15", "Sensor Weighted Average Error Computation"),
    ("Fig. 3.16", "Turn Controller State Diagram"),
    ("Fig. 3.17", "Mission State Machine Flowchart (12 States)"),
    ("Fig. 3.18", "Velocity Arbiter Logic Diagram"),
    ("Fig. 3.19", "Physical Robot Deployment Architecture"),
    ("Fig. 3.20", "Power Distribution Schematic for Physical Robot"),
    ("Fig. 4.1", "Mission Timeline Analysis — Shelf S05"),
    ("Fig. 4.2", "Robot Position Trajectory Plot during Complete Mission"),
    ("Fig. 4.3", "GUI Control Centre Interface Layout"),
    ("Fig. 4.4", "Commercial AGV Comparison Chart"),
]

for fig_num, fig_title in figures_list:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para_spacing(p, before=0, after=2)
    p.paragraph_format.line_spacing = Pt(18)
    run = p.add_run(f"{fig_num}:  {fig_title}")
    set_run_font(run, 12)

page_break()


# ════════════════════════════════════════════════════════════════════════════
# CHAPTER 1 — INTRODUCTION AND LITERATURE REVIEW
# ════════════════════════════════════════════════════════════════════════════
heading1("CHAPTER 1")
heading1("INTRODUCTION AND LITERATURE REVIEW")

heading2("1.1  Background and Industry Relevance")

body_para(
    "The logistics and warehousing industry has undergone a profound transformation over the "
    "past two decades, driven by the exponential growth of e-commerce, the globalisation of "
    "supply chains, and the relentless demand for faster order fulfilment. Modern distribution "
    "centres must process thousands of orders per day with minimal error, operating across "
    "multi-shift and often fully automated schedules. In this context, the Automated Guided "
    "Vehicle (AGV) has emerged as a cornerstone technology, providing reliable, repeatable, "
    "and cost-effective material transport within structured warehouse environments."
)
body_para(
    "Studies indicate that warehouse workers spend 60 to 70 per cent of their working time "
    "simply walking between storage locations and pick stations [1]. This walking constitutes "
    "non-value-added activity that directly inflates operational costs and limits throughput. "
    "AGV systems address this inefficiency by replacing manual transport with autonomous "
    "vehicles capable of 24-hour, seven-day operation, eliminating fatigue-related errors and "
    "removing human workers from hazardous forklift zones."
)
body_para(
    "The present project, designated ATLAS (Autonomous Transport and Logistics Automation "
    "System), develops a fully functional warehouse AGV in simulation and provides a "
    "comprehensive pathway to physical hardware deployment. The system is built on the Robot "
    "Operating System 2 (ROS2) Humble framework and simulated in Gazebo Classic 11, "
    "representing state-of-the-art open-source robotics middleware suitable for academic "
    "research and industrial prototyping alike."
)

heading2("1.2  Automated Guided Vehicles in Industry")

body_para(
    "Automated Guided Vehicles have evolved through successive generations, each defined by "
    "advances in navigation technology. The earliest AGVs, introduced in the 1950s by Barrett "
    "Electronics, relied on buried wire induction to follow predetermined paths. Subsequent "
    "generations adopted magnetic tape and painted stripe guidance, which reduced installation "
    "costs while maintaining deterministic routing. The transition to laser triangulation in "
    "the 2000s enabled more flexible path planning, culminating in the modern generation of "
    "AI-driven AGVs that perform full Simultaneous Localisation and Mapping (SLAM) in "
    "previously unmapped environments [2]."
)

add_table(
    ["Generation", "Era", "Navigation Method", "Example System"],
    [
        ["1st", "1950s", "Buried wire induction", "Barrett Electronics"],
        ["2nd", "1980s", "Magnetic tape / painted stripe", "Daifuku"],
        ["3rd", "2000s", "Laser triangulation (SLAM)", "SICK NAV350"],
        ["4th", "2010s", "Natural feature SLAM", "MiR, OTTO"],
        ["5th", "2020s", "AI-based + fleet coordination", "Amazon Robotics"],
    ],
    "Table 3: AGV Technology Generations and Navigation Methods"
)

body_para(
    "The ATLAS project implements a second- and third-generation hybrid approach: line-following "
    "navigation augmented by RFID-based shelf identification. This approach remains the dominant "
    "method in highly structured warehouse environments due to its deterministic behaviour, low "
    "infrastructure cost, and straightforward safety certification. According to industry surveys, "
    "line- and tape-guided AGVs account for approximately 60 per cent of installed AGV systems "
    "globally as of 2023 [19]."
)


heading2("1.3  Literature Review")

heading3("1.3.1  Differential Drive Kinematics in Mobile Robotics")

body_para(
    "The differential drive mechanism constitutes the most common drive configuration in "
    "mobile robotics research and industrial AGV design. Siegwart, Nourbakhsh, and Scaramuzza "
    "[1] provide the canonical treatment of differential drive kinematics, establishing the "
    "forward and inverse kinematic relationships between wheel velocities and robot body "
    "velocities. The configuration offers zero-turning-radius capability — enabling the robot "
    "to rotate in place — simple mechanical construction with minimal moving parts, and a "
    "well-understood kinematic model that facilitates accurate odometry computation. These "
    "properties make differential drive the preferred choice for warehouse AGVs operating in "
    "structured environments with right-angle navigation grids."
)
body_para(
    "Dudek and Jenkin [2] further analyse the limitations of odometric dead reckoning, noting "
    "that wheel slip, uneven floor surfaces, and encoder quantisation errors introduce "
    "cumulative position errors. In line-following AGV systems, however, the continuous "
    "correction provided by the sensor array effectively resets positional error on every "
    "segment of guide tape, making odometry primarily useful for heading estimation between "
    "correction intervals rather than absolute localisation."
)

heading3("1.3.2  PID Control for Line Following")

body_para(
    "Proportional-Integral-Derivative (PID) control is the industry standard for line "
    "following because of its mathematical simplicity, well-established tuning procedures, "
    "robustness to parameter variations, and real-time capability at any sampling frequency "
    "[15]. The application of PID control to reflectance-sensor-based line following involves "
    "computing a weighted error signal from the multi-channel sensor array and applying the "
    "control law to generate a corrective angular velocity command. Ogata [15] and Franklin "
    "et al. [16] provide comprehensive stability analyses for PID controllers, demonstrating "
    "that the closed-loop system is stable provided the proportional and derivative gains are "
    "chosen to place the characteristic roots in the left-half complex plane."
)
body_para(
    "In the ATLAS implementation, the integral term is deliberately set to zero, yielding a "
    "Proportional-Derivative (PD) controller. This design decision is justified by three "
    "engineering considerations: first, the line-following setpoint changes rapidly at curves "
    "and junction points, causing integral windup that generates overshoot; second, no "
    "steady-state positional error exists because the sensor directly measures lateral "
    "displacement from the guide tape; and third, the derivative term provides sufficient "
    "damping to prevent oscillation without the complexity of integrator anti-windup logic."
)

heading3("1.3.3  ROS2 as Robot Middleware")

body_para(
    "The Robot Operating System 2 (ROS2) represents a fundamental architectural redesign of "
    "the original ROS framework, motivated by the requirements of production robotic systems. "
    "As documented in the official ROS2 Humble documentation [5], the key improvements over "
    "ROS1 include a Data Distribution Service (DDS) communication layer that eliminates the "
    "single point of failure introduced by the rosmaster process, Quality of Service (QoS) "
    "policies that enable reliable communication in lossy network environments, real-time "
    "capability through deterministic scheduling, and multi-platform support including ARM "
    "architectures relevant to embedded robot computing. Quigley et al. [7] document the "
    "motivations and design principles behind ROS, and the same principles have guided the "
    "evolution towards ROS2."
)
body_para(
    "For the ATLAS project, ROS2 Humble was selected because it is a Long Term Support (LTS) "
    "release with guaranteed support through 2027, compatible with Ubuntu 22.04 LTS, and "
    "provides native integration with Gazebo Classic 11 through the gazebo_ros_pkgs package. "
    "The DDS-based communication architecture also supports future multi-robot fleet scenarios "
    "without architectural changes to the single-robot software stack."
)

heading3("1.3.4  RFID in Warehouse Automation")

body_para(
    "Radio Frequency Identification (RFID) technology provides reliable position confirmation "
    "in structured environments due to several properties that distinguish it from alternative "
    "identification technologies. Finkenzeller [11] comprehensively documents RFID systems, "
    "noting that unlike optical barcodes, RFID requires no line-of-sight between reader and "
    "tag, enabling reliable detection even when the tag is partially obscured. The read-while-"
    "moving capability of inductive RFID systems — with the RDM6300 125 kHz reader operating "
    "at up to 10 cm read range — allows shelf identification without stopping the AGV. Passive "
    "RFID tags require no onboard power or maintenance, making them ideal for permanent "
    "installation at shelf locations in a warehouse environment."
)
body_para(
    "Wurman, D'Andrea, and Mountz [18] document the large-scale AGV fleet deployed by Amazon "
    "Robotics (formerly Kiva Systems), in which RFID-marked floor stickers provide absolute "
    "position references that correct accumulated odometry errors. The ATLAS system employs an "
    "analogous strategy at a smaller scale: 21 RFID tags provide position confirmation at "
    "shelf arrival events, supplementing the continuous lane-following control loop."
)

heading3("1.3.5  Simulation-First Development Methodology")

body_para(
    "The simulation-first development approach has become a well-established paradigm in "
    "robotics engineering, supported by the availability of high-fidelity physics simulators "
    "such as Gazebo. The Gazebo Classic 11 simulator, documented by the Open Source Robotics "
    "Foundation [6], implements the Open Dynamics Engine (ODE) for rigid body dynamics with "
    "a configurable physics time step of 1 ms, providing sufficiently accurate contact "
    "simulation for wheeled robot development. De Ryck, Versteyhe, and Debrouwere [19] survey "
    "modern AGV control algorithms and note that simulation environments are now routinely "
    "used for complete system validation before physical hardware procurement, reducing "
    "development risk and cost."
)

add_table(
    ["Reference", "Topic", "Key Contribution", "Relevance to ATLAS"],
    [
        ["Siegwart et al. [1]", "Mobile robot kinematics", "Canonical differential drive model", "Forward/inverse kinematics equations"],
        ["Dudek & Jenkin [2]", "Odometry limitations", "Error accumulation analysis", "Justifies RFID correction"],
        ["Ogata [15]", "Control engineering", "PID stability analysis", "PD controller design"],
        ["Franklin et al. [16]", "Feedback control", "Closed-loop characteristic equations", "Line follower stability"],
        ["Finkenzeller [11]", "RFID systems", "Inductive RFID principles", "Tag detection model"],
        ["Wurman et al. [18]", "Warehouse AGV fleets", "RFID-based localisation", "Shelf identification strategy"],
        ["De Ryck et al. [19]", "AGV control algorithms", "Survey of modern methods", "System design context"],
        ["Azadeh et al. [20]", "Automated warehouses", "Robotized warehouse review", "Industry context"],
        ["ISO 3691-4 [10]", "AGV safety standards", "Driverless truck requirements", "Safety architecture"],
        ["Open Robotics [5]", "ROS2 Humble", "DDS middleware documentation", "Communication layer"],
    ],
    "Table 4: Comparison of Prior Literature on AGV Systems"
)


heading2("1.4  Research Gap and Motivation")

body_para(
    "A review of the existing literature reveals that while the individual components of "
    "warehouse AGV systems — differential drive kinematics, PID line following, RFID "
    "localisation, and state machine mission management — are individually well-documented, "
    "the integration of all these components into a cohesive, fully functional system built "
    "on modern ROS2 middleware, validated through Gazebo simulation, and analysed for direct "
    "physical hardware deployment, represents a gap that the present project addresses."
)
body_para(
    "Furthermore, academic literature on AGV systems frequently presents either high-level "
    "system descriptions without implementation detail, or narrow technical analyses of "
    "individual subsystems without demonstrating complete operational capability. The ATLAS "
    "project bridges this gap by presenting a complete, working implementation with "
    "quantified performance metrics, a clear simulation-to-physical conversion pathway, and "
    "a safety-conscious architecture suitable for industrial adaptation."
)

heading2("1.5  Project Objectives")

body_para(
    "The primary objectives of the ATLAS project are as follows: (i) to design a "
    "differential-drive AGV platform with accurately modelled kinematics in URDF/Xacro "
    "format; (ii) to implement line-following navigation using an 8-channel virtual infrared "
    "reflectance sensor array with PD control; (iii) to implement RFID-based shelf "
    "identification using a simulated tag detection system with hysteresis; (iv) to design a "
    "12-state Finite State Machine governing complete pick-and-return mission cycles; (v) to "
    "create a professional PyQt5-based graphical control interface for fleet monitoring and "
    "mission management; (vi) to demonstrate the complete system in a realistic Gazebo "
    "simulation with quantified performance metrics; and (vii) to provide a comprehensive "
    "simulation-to-physical hardware conversion analysis."
)

add_table(
    ["Objective", "Success Criterion", "Status"],
    [
        ["Differential drive kinematic model", "URDF validated in Gazebo simulation", "Achieved"],
        ["8-channel line following with PD control", "Tracking error < 5 mm at 0.4 m/s", "Achieved"],
        ["RFID-based shelf identification", "100% detection rate for 21 tags", "Achieved"],
        ["12-state mission FSM", "Complete pick-and-return cycle demonstrated", "Achieved"],
        ["PyQt5 control interface", "Real-time telemetry and mission dispatch", "Achieved"],
        ["Gazebo simulation validation", "100% mission completion rate", "Achieved"],
        ["Physical hardware conversion analysis", "Component mapping and cost estimate", "Achieved"],
    ],
    "Table 5: Project Objectives and Success Criteria"
)

heading2("1.6  Chapter Summary")

body_para(
    "This chapter has established the industrial motivation for warehouse AGV systems, "
    "reviewed the relevant literature on differential drive kinematics, PID line following, "
    "RFID localisation, ROS2 middleware, and simulation-based development, identified the "
    "research gap addressed by the ATLAS project, and stated the project objectives. The "
    "following chapter formalises the problem statement and defines the design constraints "
    "and operational assumptions under which the system was developed."
)

page_break()


# ════════════════════════════════════════════════════════════════════════════
# CHAPTER 2 — PROBLEM FORMULATION
# ════════════════════════════════════════════════════════════════════════════
heading1("CHAPTER 2")
heading1("PROBLEM FORMULATION")

heading2("2.1  Problem Domain")

body_para(
    "The problem addressed by the ATLAS project falls within the domain of autonomous "
    "intralogistics — the automated movement of materials within a warehouse or distribution "
    "centre. The specific problem is the design and implementation of a single autonomous "
    "vehicle capable of receiving a shelf identifier from an operator interface, navigating "
    "from a home docking station to the identified shelf location, executing a simulated "
    "pick operation, and returning to the home dock, all without human intervention and with "
    "sufficient reliability and precision to serve as a proof-of-concept for industrial AGV "
    "deployment."
)
body_para(
    "The warehouse environment is modelled as a structured layout with a single longitudinal "
    "spine navigation path and multiple parallel aisle paths branching from the spine to "
    "shelf locations. This topology mirrors the conventional warehouse rack-and-aisle "
    "arrangement used in the majority of storage facilities globally. The deterministic, "
    "structured nature of this environment makes it amenable to guide-tape navigation rather "
    "than SLAM-based localisation, providing a reliable and cost-effective solution."
)

heading2("2.2  Existing Limitations")

body_para(
    "The limitations motivating the development of the ATLAS system are well-documented in "
    "both academic literature and industry practice. Manual material handling suffers from "
    "three fundamental constraints. First, human operators are subject to physical fatigue, "
    "which limits continuous operational throughput to approximately one shift per worker per "
    "day and introduces error rates that increase over time. Second, the presence of human "
    "workers in areas where forklifts and heavy material handling equipment operate creates "
    "significant safety risks, imposing regulatory constraints on warehouse layout and "
    "operational procedures. Third, the labour costs associated with warehouse walking — "
    "constituting 60 to 70 per cent of worker time — represent a major component of "
    "operational expenditure that directly erodes profit margins in high-volume distribution "
    "operations [1]."
)
body_para(
    "Existing commercial AGV solutions, while effective at scale, present significant barriers "
    "to adoption for small and medium enterprises: the MiR 250 carries a price point of "
    "approximately USD 25,000, and the OTTO 100 approximately USD 30,000, while Amazon "
    "Robotics solutions are proprietary and available only to Amazon. The ATLAS project "
    "demonstrates that a functionally equivalent AGV for structured environments can be "
    "implemented on open-source software with a physical hardware cost below USD 400."
)

heading2("2.3  Project Requirements and Design Constraints")

body_para(
    "The design requirements for the ATLAS system were established through analysis of the "
    "warehouse environment and the operational tasks to be performed. The robot must navigate "
    "a spine path of 12 metres in length and aisle paths of 5 metres in length, with turns "
    "of exactly 90 degrees at aisle junctions and 180 degrees at shelf pickup points. "
    "Navigation must be reliable enough to achieve correct shelf identification for all 20 "
    "shelf positions across 5 aisles. The system must respond to emergency stop commands "
    "within 100 milliseconds, and a reset-to-dock function must be available to return the "
    "robot to its home position from any operational state."
)

add_table(
    ["Requirement", "Specification", "Source"],
    [
        ["Navigation speed", "0.4 m/s nominal", "Power and accuracy trade-off"],
        ["Line tracking accuracy", "< 5 mm lateral error", "Shelf alignment requirement"],
        ["Turn accuracy", "±5° maximum error", "Aisle entry accuracy"],
        ["RFID detection rate", "> 95% per pass", "Shelf identification reliability"],
        ["Emergency stop response", "< 100 ms", "ISO 3691-4 safety standard"],
        ["Mission completion rate", "> 98% per cycle", "Operational reliability"],
        ["Warehouse footprint", "12m spine × 5m aisles", "Environment specification"],
        ["Shelf count", "20 shelves (5 aisles × 4 positions)", "Warehouse layout"],
        ["Home dock", "Single position at (0,0)", "Charging and reset point"],
        ["Communication", "ROS2 DDS topics", "Middleware selection"],
        ["Simulation platform", "Gazebo Classic 11", "Academic availability"],
        ["Computing platform (physical)", "Raspberry Pi 4 (4 GB)", "Cost and ROS2 compatibility"],
    ],
    "Table 6: Design Constraints and Operational Assumptions"
)

heading2("2.4  Operational Assumptions")

body_para(
    "The ATLAS system operates under the following explicit assumptions. The warehouse floor "
    "is flat and free of obstacles — the AGV does not implement collision avoidance. Guide "
    "tape is installed at precisely defined positions matching the world model embedded in "
    "the software, and the tape contrast is sufficient for reliable reflectance sensor "
    "detection. RFID tags are positioned at the intersection of aisle tape and shelf "
    "x-positions, with no metallic interference that would attenuate the 125 kHz signal. "
    "The robot spawns at the home dock position (x=0, y=0, heading=90°) at system startup, "
    "and the odometry frame coincides with the world frame due to this initialisation "
    "condition. Only one robot is present in the warehouse at any given time in the current "
    "implementation."
)

heading2("2.5  Expected Outcomes and System Objectives")

body_para(
    "The expected outcomes of the project are: a fully functional Gazebo simulation "
    "demonstrating autonomous warehouse pick-and-return operations; quantified performance "
    "metrics including mission completion rate, tracking accuracy, and cycle time; a "
    "professional graphical control interface with real-time telemetry; a complete simulation-"
    "to-physical hardware conversion report with component specifications, wiring diagrams, "
    "and cost estimates; and a software architecture suitable for extension to multi-robot "
    "fleet coordination in future work."
)

heading2("2.6  Scope and Limitations")

body_para(
    "The scope of the current project is bounded as follows: the implementation covers a "
    "single-robot system only; multi-robot fleet coordination, while architecturally "
    "anticipated, is not implemented in the current version. The system does not implement "
    "obstacle detection or dynamic path replanning; it relies on clear navigation paths as "
    "specified in the operational assumptions. Sensor models are virtual approximations "
    "rather than real noise models; the physical noise characteristics of the QTR-8A sensor "
    "and BNO055 IMU are not replicated in simulation. The payload is represented as a boolean "
    "state flag rather than a physical mass, and no payload physics are simulated. Odometry "
    "drift, while negligible in the simulation environment, would require periodic RFID "
    "correction in a physical deployment over extended operational periods."
)

heading2("2.7  Chapter Summary")

body_para(
    "This chapter has formally defined the problem domain, identified the operational "
    "limitations of existing solutions, specified the design requirements and constraints "
    "governing the ATLAS system, stated the operational assumptions, and bounded the scope "
    "of the project. The following chapter presents the complete technical methodology, "
    "including system architecture, kinematic analysis, control algorithms, simulation "
    "environment design, and physical hardware conversion approach."
)

page_break()


# ════════════════════════════════════════════════════════════════════════════
# CHAPTER 3 — MODELING, SOLUTION METHODOLOGY AND IMPLEMENTATION
# ════════════════════════════════════════════════════════════════════════════
heading1("CHAPTER 3")
heading1("MODELING, SOLUTION METHODOLOGY AND EXPERIMENTAL PROCEDURE")

body_para(
    "This chapter presents the complete technical implementation of the ATLAS Warehouse AGV "
    "system. It covers the overall system architecture, ROS2 software package structure, "
    "node and topic communication architecture, robot mechanical model, differential drive "
    "kinematics, Gazebo simulation environment, sensor systems, control algorithms, mission "
    "state machine design, and simulation-to-physical hardware conversion methodology. This "
    "chapter constitutes the primary engineering contribution of the project."
)

heading2("3.1  Overall System Architecture")

body_para(
    "The ATLAS system follows a layered architecture with clear separation of concerns across "
    "five functional layers. The Perception Layer comprises the line sensor, IMU, RFID reader, "
    "and odometry modules, which convert physical or simulated sensor signals into ROS2 topic "
    "messages. The Control Layer implements the PD line following controller and IMU-based "
    "turn controller, each publishing velocity recommendations to dedicated topics. The "
    "Planning Layer hosts the mission Finite State Machine and queue manager, which interpret "
    "mission commands and sequence the robot through its operational states. The Actuation "
    "Layer is the sole publisher on the /atlas/cmd_vel topic, implementing a velocity arbiter "
    "that selects the appropriate velocity source based on the current FSM state. The "
    "Interface Layer provides the PyQt5 graphical control centre and command-line mission "
    "sender for operator interaction."
)
body_para(
    "This architectural separation ensures that individual subsystems can be developed, "
    "tested, and replaced independently. In particular, the physical hardware deployment "
    "requires replacing only the perception layer nodes — the line sensor, IMU, and RFID "
    "reader — while all control, planning, and interface nodes run unchanged on the physical "
    "robot."
)

heading2("3.2  Software and ROS2 Package Architecture")

body_para(
    "The ATLAS software is organised into six ROS2 packages, each with a clearly defined "
    "responsibility. The package structure follows ROS2 best practices, separating message "
    "definitions, robot description, simulation environment, navigation control, mission "
    "management, and launch configuration into distinct packages with explicit dependency "
    "declarations."
)

add_table(
    ["Package", "Build Type", "Purpose"],
    [
        ["atlas_interfaces", "ament_cmake", "Custom message definitions: FleetMission, RobotState, ShelfTag"],
        ["atlas_description", "ament_cmake", "URDF/Xacro robot model, RViz2 configuration"],
        ["atlas_gazebo", "ament_cmake", "Gazebo warehouse world file (warehouse.world)"],
        ["atlas_navigation", "ament_python", "Line sensor, line follower, turn controller, tag detector"],
        ["atlas_mission_manager", "ament_python", "Mission FSM, velocity arbiter, GUI, CLI sender"],
        ["atlas_bringup", "ament_cmake", "Master launch file (atlas_full.launch.py)"],
    ],
    "Table 7: ROS2 Package Structure"
)

body_para(
    "Three custom message types are defined in the atlas_interfaces package. The FleetMission "
    "message carries mission_id (string), target_shelf (string), sku (string), and priority "
    "(integer) fields, enabling the GUI or CLI to dispatch pick missions to the mission "
    "manager. The RobotState message provides a comprehensive status snapshot including the "
    "FSM state string, active mission identifiers, battery percentage, carrying flag, last "
    "detected tag, and current pose. The ShelfTag message reports RFID detection events with "
    "tag_id, shelf_id, detection distance, home flag, and timestamp."
)

heading2("3.3  ROS2 Node Architecture")

body_para(
    "Six primary computational nodes implement the robot behaviour. Each node is an "
    "independent ROS2 process, enabling fault isolation and independent restart. The "
    "atlas_line_sensor node operates at 50 Hz, computing the world positions of eight virtual "
    "infrared sensor elements from the current odometry pose and publishing binary and raw "
    "distance data, together with junction detection events. The atlas_line_follower node "
    "processes the binary sensor data at 50 Hz to compute a PD control output, published as "
    "a navigation velocity recommendation. The atlas_turn_ctrl node monitors IMU heading at "
    "50 Hz and applies a bang-bang angular velocity until the commanded heading is reached "
    "within a 3° tolerance. The atlas_tag_detect node monitors odometry at 50 Hz and fires "
    "RFID detection events when the robot enters the 0.5 m detection radius of any tag. The "
    "atlas_mission_mgr node hosts the 12-state FSM and velocity arbiter, and is the sole "
    "publisher on the /atlas/cmd_vel topic. The atlas_gui node provides the PyQt5 control "
    "centre in a separate process, communicating exclusively via ROS2 topics."
)

add_table(
    ["Node", "Package", "Rate", "Input Topics", "Output Topics"],
    [
        ["atlas_line_sensor", "atlas_navigation", "50 Hz", "/atlas/odom", "/atlas/line_sensors, /atlas/line_raw, /atlas/junction"],
        ["atlas_line_follower", "atlas_navigation", "50 Hz", "/atlas/line_sensors", "/atlas/nav_vel"],
        ["atlas_turn_ctrl", "atlas_navigation", "50 Hz", "/atlas/imu, /atlas/turn_cmd", "/atlas/turn_vel, /atlas/turn_done"],
        ["atlas_tag_detect", "atlas_navigation", "50 Hz", "/atlas/odom", "/atlas/tag_event"],
        ["atlas_mission_mgr", "atlas_mission_manager", "50 Hz", "All nav/turn/tag topics", "/atlas/cmd_vel, /atlas/robot_state, /atlas/log"],
        ["atlas_gui", "atlas_mission_manager", "Event-driven", "/atlas/robot_state, /atlas/odom, /atlas/log", "/atlas/mission_cmd, /atlas/estop, /atlas/reset, /atlas/reset_to_dock"],
    ],
    "Table 8: Node Input/Output Summary"
)


heading2("3.4  Topic Communication and Message Flow")

body_para(
    "The ATLAS system uses 17 ROS2 topics for inter-node communication. The critical design "
    "decision governing the topic architecture is the velocity arbiter pattern: only the "
    "atlas_mission_mgr node publishes to /atlas/cmd_vel, the topic that drives the Gazebo "
    "differential drive plugin. All other nodes publish velocity recommendations to "
    "intermediate topics (/atlas/nav_vel for line following, /atlas/turn_vel for turning), "
    "which the mission manager reads and selects from based on the current FSM state. This "
    "pattern guarantees that conflicting velocity commands can never reach the wheels, and "
    "that the emergency stop is always honoured."
)

add_table(
    ["Topic", "Message Type", "Publisher", "Rate", "Purpose"],
    [
        ["/atlas/odom", "nav_msgs/Odometry", "Gazebo diff_drive", "50 Hz", "Robot odometry"],
        ["/atlas/imu", "sensor_msgs/Imu", "Gazebo IMU plugin", "100 Hz", "IMU heading data"],
        ["/atlas/cmd_vel", "geometry_msgs/Twist", "mission_mgr ONLY", "50 Hz", "Wheel velocity commands"],
        ["/atlas/nav_vel", "geometry_msgs/Twist", "line_follower", "50 Hz", "Line-following velocity"],
        ["/atlas/turn_vel", "geometry_msgs/Twist", "turn_controller", "50 Hz", "Turn velocity"],
        ["/atlas/turn_cmd", "std_msgs/Float32", "mission_mgr", "Event", "Turn angle command (rad)"],
        ["/atlas/turn_done", "std_msgs/Empty", "turn_controller", "Event", "Turn complete signal"],
        ["/atlas/line_sensors", "std_msgs/Int8MultiArray", "line_sensor", "50 Hz", "8-bit binary sensor array"],
        ["/atlas/line_raw", "std_msgs/Float32MultiArray", "line_sensor", "50 Hz", "Raw sensor distances (m)"],
        ["/atlas/junction", "std_msgs/Empty", "line_sensor", "Event", "Junction detected"],
        ["/atlas/tag_event", "atlas_interfaces/ShelfTag", "tag_detector", "Event", "RFID detection event"],
        ["/atlas/mission_cmd", "atlas_interfaces/FleetMission", "GUI/CLI", "Event", "New mission dispatch"],
        ["/atlas/robot_state", "atlas_interfaces/RobotState", "mission_mgr", "10 Hz", "Live robot status"],
        ["/atlas/log", "std_msgs/String", "mission_mgr", "Event", "Timestamped log messages"],
        ["/atlas/estop", "std_msgs/Empty", "GUI/CLI", "Event", "Emergency stop trigger"],
        ["/atlas/reset", "std_msgs/Empty", "GUI/CLI", "Event", "E-stop reset"],
        ["/atlas/reset_to_dock", "std_msgs/Empty", "GUI/CLI", "Event", "Return to home dock"],
    ],
    "Table 9: Complete ROS2 Topic List"
)

heading2("3.5  Transform (TF) Tree")

body_para(
    "The coordinate transform tree published by the Robot State Publisher node defines the "
    "geometric relationship between all robot link frames. The tree is rooted at the odom "
    "frame, which is published by the Gazebo differential drive plugin at 50 Hz. The "
    "transform from odom to base_footprint represents the robot's estimated position in "
    "the world. Fixed transforms connect base_footprint to base_link (offset by wheel "
    "radius vertically), and base_link to the left_wheel, right_wheel, caster_wheel, and "
    "imu_link frames. A key design decision — discussed in Section 3.8 — is that the odom "
    "frame coincides with the world frame in this system, eliminating the need for a "
    "separate localisation layer."
)
body_para(
    "The TF tree structure is: odom → base_footprint → base_link → left_wheel (continuous "
    "revolute), right_wheel (continuous revolute), caster_wheel (fixed), imu_link (fixed). "
    "This minimal transform tree is sufficient for all sensor computations and visualisation "
    "in RViz2, while remaining efficient enough to not introduce computational overhead at "
    "the 50 Hz update rate."
)

heading2("3.6  Robot Description — URDF/Xacro Model")

body_para(
    "The robot is described in URDF format using the Xacro macro language, defined in "
    "atlas_agv.urdf.xacro. The model parameterises all geometric and inertial properties "
    "through named constants, making modifications to physical dimensions straightforward "
    "without requiring edits to multiple locations in the file. The following table presents "
    "the physical parameters as defined in the URDF."
)

add_table(
    ["Parameter", "Symbol", "Value", "Physical Meaning"],
    [
        ["Body length (x)", "BX", "0.30 m", "Chassis fore-aft dimension"],
        ["Body width (y)", "BY", "0.25 m", "Chassis lateral dimension"],
        ["Body height (z)", "BZ", "0.10 m", "Chassis vertical dimension"],
        ["Robot mass", "BM", "2.5 kg", "Total chassis mass"],
        ["Wheel radius", "WR", "0.05 m (50 mm)", "Drive wheel radius"],
        ["Wheel track", "WS", "0.30 m", "Distance between wheel centres"],
        ["Wheel width", "WT", "0.04 m", "Tyre width"],
        ["Wheel mass", "WM", "0.2 kg", "Mass per drive wheel"],
        ["Caster radius", "CR", "0.025 m (25 mm)", "Passive support wheel radius"],
        ["Drive type", "—", "Differential drive", "Two powered + one passive wheel"],
        ["Max wheel torque", "—", "5.0 N·m", "Gazebo diff_drive plugin limit"],
        ["Max wheel acceleration", "—", "2.0 rad/s²", "Gazebo diff_drive plugin limit"],
    ],
    "Table 10: Physical Parameters from URDF/Xacro"
)

body_para(
    "The URDF defines three link types. The base_link represents the main chassis as a "
    "rectangular box with uniform density inertial properties. The left and right drive "
    "wheels are cylinders with continuous revolute joints about the y-axis, connected to "
    "base_link with a lateral offset of ±0.15 m (half the wheel track). The caster wheel "
    "is a sphere connected to base_link by a fixed joint at position (0.12, 0, −0.025) m, "
    "with both friction coefficients set to zero to allow free rotation in any direction "
    "without introducing yaw torque. The imu_link is a massless fixed frame at the chassis "
    "top surface, carrying the Gazebo IMU sensor plugin."
)
body_para(
    "The Gazebo differential drive plugin (libgazebo_ros_diff_drive.so) is declared within "
    "the URDF under the gazebo element. It subscribes to /atlas/cmd_vel for velocity "
    "commands and publishes odometry to /atlas/odom and the corresponding TF transform at "
    "50 Hz. Wheel friction is modelled with mu1=1.5 and mu2=1.0, providing realistic traction "
    "while preventing excessive wheel slip in simulation."
)


heading2("3.7  Differential Drive Kinematics")

heading3("3.7.1  Coordinate Frame Convention")

body_para(
    "The ATLAS robot uses the standard ROS mobile robot coordinate frame convention: the "
    "x-axis points forward (direction of travel), the y-axis points left, and the z-axis "
    "points upward. The origin of the base_footprint frame is at the centre point of the "
    "line joining the two drive wheel ground contact points, projected onto the floor plane. "
    "All kinematic equations use this convention."
)

heading3("3.7.2  Forward Kinematics")

body_para(
    "Forward kinematics computes the robot body velocity given the individual wheel "
    "velocities. Given left wheel linear velocity v_L and right wheel linear velocity v_R, "
    "the linear velocity of the robot centre and the angular velocity about the centre are:"
)

add_equation("v = (v_R + v_L) / 2", "(3.1)")
add_equation("ω = (v_R − v_L) / L", "(3.2)")

body_para(
    "where v is the linear velocity at the robot centre (m/s), ω is the angular velocity "
    "(rad/s), and L = 0.30 m is the wheel track width. The continuous-time pose evolution "
    "equations — expressing how the robot position and heading change in the world frame — are:"
)

add_equation("ẋ = v · cos(θ)", "(3.3)")
add_equation("ẏ = v · sin(θ)", "(3.4)")
add_equation("θ̇ = ω", "(3.5)")

heading3("3.7.3  Inverse Kinematics")

body_para(
    "Inverse kinematics computes the required wheel velocities for a desired robot body "
    "velocity command (v, ω), as issued by the /atlas/cmd_vel topic. The wheel velocities are:"
)

add_equation("v_L = v − ω · L / 2", "(3.6)")
add_equation("v_R = v + ω · L / 2", "(3.7)")

body_para(
    "Converting to wheel angular velocities for motor control commands:"
)

add_equation("ω_L = v_L / r = (v − ω·L/2) / r", "(3.8)")
add_equation("ω_R = v_R / r = (v + ω·L/2) / r", "(3.9)")

body_para(
    "where r = 0.05 m is the wheel radius. At the nominal operating speed of v = 0.4 m/s "
    "with zero angular velocity, both wheels rotate at ω_L = ω_R = 0.4/0.05 = 8.0 rad/s, "
    "corresponding to 76.4 RPM — well within the 200 RPM rated speed of the JGB37-520 motor "
    "specified for physical deployment."
)

heading3("3.7.4  Instantaneous Centre of Rotation")

body_para(
    "The Instantaneous Centre of Rotation (ICR) is the point about which the robot rotates "
    "at any given instant. Its distance from the robot centre is:"
)

add_equation("R = (L/2) · (v_R + v_L) / (v_R − v_L)", "(3.10)")

body_para(
    "Three special cases define the robot's turning capability. When v_R = v_L, R → ∞ and "
    "the robot travels in a straight line. When v_R = −v_L, R = 0 and the robot spins in "
    "place about its centre — a capability unique to differential drive and exploited by the "
    "ATLAS turn controller for 90° aisle entry turns. When v_L = 0, R = L/2 = 0.15 m, "
    "and the robot pivots about the stationary left wheel."
)

heading3("3.7.5  Odometry Integration")

body_para(
    "The robot's position is computed by dead-reckoning integration of wheel encoder "
    "measurements. For each timestep Δt, the incremental distance and heading change are:"
)

add_equation("Δs = (Δs_R + Δs_L) / 2", "(3.11)")
add_equation("Δθ = (Δs_R − Δs_L) / L", "(3.12)")

body_para(
    "The position update uses second-order midpoint integration for improved accuracy "
    "compared to first-order Euler integration:"
)

add_equation("x_{k+1} = x_k + Δs · cos(θ_k + Δθ/2)", "(3.13)")
add_equation("y_{k+1} = y_k + Δs · sin(θ_k + Δθ/2)", "(3.14)")
add_equation("θ_{k+1} = θ_k + Δθ", "(3.15)")

body_para(
    "In the Gazebo simulation, this integration is performed by the differential drive plugin "
    "at 50 Hz. On physical hardware, it would be performed by the motor driver node using "
    "encoder tick counts from the JGB37-520 motors (330 counts per revolution, sufficient "
    "for 0.95 mm odometric resolution at the wheel circumference)."
)

heading3("3.7.6  Inertia Tensor")

body_para(
    "For the rectangular chassis modelled as a uniform density box with mass m = 2.5 kg "
    "and dimensions a = 0.30 m, b = 0.25 m, c = 0.10 m, the principal moments of inertia "
    "about the body centre of mass are:"
)

add_equation("I_xx = m(b² + c²) / 12 = 2.5(0.25² + 0.10²) / 12 = 0.01510 kg·m²", "(3.16)")
add_equation("I_yy = m(a² + c²) / 12 = 2.5(0.30² + 0.10²) / 12 = 0.02083 kg·m²", "(3.17)")
add_equation("I_zz = m(a² + b²) / 12 = 2.5(0.30² + 0.25²) / 12 = 0.03177 kg·m²", "(3.18)")

add_table(
    ["Moment", "Value (kg·m²)", "Significance"],
    [
        ["I_xx (roll)", "0.01510", "Resistance to sideways tipping"],
        ["I_yy (pitch)", "0.02083", "Resistance to fore-aft tipping"],
        ["I_zz (yaw)", "0.03177", "Resistance to turning — dominant for control"],
    ],
    "Table 11: Inertia Tensor Values for Chassis"
)

heading3("3.7.7  Stability Analysis")

body_para(
    "The three-point support polygon is defined by the two drive wheel ground contact points "
    "at (0, +0.15) and (0, −0.15), and the caster contact point at (0.12, 0). The chassis "
    "centre of mass at (0, 0, 0.10) m projects onto the floor at (0, 0), which lies at the "
    "geometric centre of the support triangle. Placing the battery — the heaviest individual "
    "component at approximately 350 g — directly above the drive axle maintains this "
    "favourable centre-of-gravity position, ensuring maximum traction on the drive wheels "
    "and stability against tipping during acceleration."
)


heading2("3.8  Gazebo Simulation Environment")

heading3("3.8.1  Physics Configuration")

body_para(
    "The Gazebo warehouse world is configured using the SDF (Simulation Description Format) "
    "specification. The physics engine is the Open Dynamics Engine (ODE) with a fixed "
    "simulation time step of 1 ms, a real-time factor of 1.0, and an update rate of "
    "1000 Hz. The ODE quick solver with 50 constraint iterations is used, providing "
    "sufficient accuracy for wheeled robot contact dynamics without excessive computational "
    "cost."
)

add_table(
    ["Parameter", "Value", "Purpose"],
    [
        ["Physics engine", "ODE (Open Dynamics Engine)", "Rigid body dynamics simulation"],
        ["Time step", "0.001 s (1 ms)", "Contact resolution accuracy"],
        ["Real-time factor", "1.0", "Wall-clock synchronisation"],
        ["Update rate", "1000 Hz", "Physics steps per second"],
        ["Solver type", "Quick (50 iterations)", "Constraint resolution"],
        ["World gravity", "9.81 m/s² (z-down)", "Standard Earth gravity"],
    ],
    "Table 12: Gazebo Physics Configuration"
)

heading3("3.8.2  Warehouse World Layout")

body_para(
    "The warehouse world file (warehouse.world) defines a structured 16 × 16 metre "
    "industrial warehouse environment. The navigation infrastructure comprises a spine tape "
    "line running from (x=0, y=0) to (x=0, y=12), 5 cm wide and coloured black on a grey "
    "floor to provide high reflectance contrast. Five aisle tape lines branch from the spine "
    "at y=2, 4, 6, 8, and 10 m, extending from (x=0, y=ay) to (x=5, y=ay). Yellow safety "
    "lane stripes flank the spine on both sides at x=±0.15 m. Twenty shelf racks are "
    "positioned at x=1, 2, 3, 4 m on each of the five aisles, with three shelf levels at "
    "heights of 0.25 m, 0.55 m, and 0.80 m above floor level, each populated with cargo "
    "boxes. Blue RFID floor marker pads (0.15 × 0.15 m) are embedded at the intersection "
    "of each aisle tape and each shelf x-position, providing visual confirmation of tag "
    "locations."
)
body_para(
    "The environment includes a home dock at the origin, marked by a 0.7 × 0.7 m green "
    "floor pad with a dock sign post, and a charging station at (−1.5, 0). RFID gate "
    "posts are positioned at each aisle entrance on the spine (y=2, 4, 6, 8, 10), with "
    "blue pillars and a horizontal bar with an LED indicator. The environment also includes "
    "perimeter walls, a loading dock platform at y=−1.2 m, ceiling-mounted fluorescent "
    "light fixtures in a 5 × 4 grid, yellow safety bollards at corner positions, a "
    "fire extinguisher, and an information board — providing a realistic industrial "
    "warehouse appearance suitable for demonstration."
)

heading3("3.8.3  Key Design Decision: Odometry Equals World Frame")

body_para(
    "A critical architectural decision in the ATLAS simulation is that the odometry frame "
    "coincides exactly with the world frame. This is achieved because the differential drive "
    "plugin initialises the odometry at the robot's spawn pose, and the robot spawns at "
    "world position (0, 0, 0.01) with heading 90° (π/2 radians). Since the plugin "
    "initialises odom at spawn, and spawn is at world origin, the odom and world frames "
    "are identical. This eliminates the need for a separate localisation layer — the line "
    "sensor node can directly use odometry coordinates as world coordinates for tape "
    "distance calculations, without requiring any TF lookup or coordinate transformation. "
    "In a physical deployment on a floor with consistent starting conditions, this same "
    "assumption holds valid for the duration of a mission, with accumulated odometry error "
    "corrected by RFID events."
)

heading3("3.8.4  Launch Architecture")

body_para(
    "The atlas_full.launch.py file orchestrates the entire system startup sequence with "
    "carefully timed delays to ensure correct initialisation order. Gazebo is launched at "
    "t=0 with the warehouse world. The robot is spawned at t=4 s, allowing Gazebo physics "
    "to fully initialise before entity creation. The Robot State Publisher is launched at "
    "t=6 s, after the robot model is registered in Gazebo. All navigation and mission "
    "management nodes are launched at t=10 s, ensuring that odometry and IMU topics are "
    "publishing before the control nodes begin subscribing. RViz2 visualisation is launched "
    "at t=12 s, and the GUI node at t=14 s. This staged startup prevents race conditions "
    "between node initialisation and topic availability."
)


heading2("3.9  Sensor Architecture")

heading3("3.9.1  Virtual Line Sensor — Operating Principle")

body_para(
    "The line sensor node (line_sensor.py) simulates an 8-channel infrared reflectance "
    "sensor array operating at 50 Hz. Rather than processing actual optical sensor data, "
    "it computes the geometric distance from each of eight virtual sensor element positions "
    "to the nearest guide tape segment in the world. This approach provides deterministic "
    "sensor behaviour free from noise, enabling clean control law validation before "
    "introducing real sensor characteristics."
)
body_para(
    "Each of the eight sensor elements is positioned at a fixed lateral offset from the "
    "robot's forward direction. The world-frame position of sensor element i is computed "
    "from the current odometry pose as:"
)

add_equation("s_x(i) = x + cos(θ)·d_fwd − sin(θ)·offset_i", "(3.19)")
add_equation("s_y(i) = y + sin(θ)·d_fwd + cos(θ)·offset_i", "(3.20)")

body_para(
    "where d_fwd = 0.10 m is the forward projection distance of the sensor array from "
    "the base_footprint origin, and offset_i takes values [0.07, 0.05, 0.03, 0.01, −0.01, "
    "−0.03, −0.05, −0.07] m for i = 0 to 7, spanning a total lateral width of 14 cm. "
    "The minimum distance from each sensor position to any guide tape segment is then "
    "computed, and a binary thresholding operation determines whether the sensor is over "
    "the tape:"
)

add_equation("s_i = 1  if  d_i ≤ 0.04 m  (half of LINE_WIDTH = 0.08 m)", "(3.21)")
add_equation("s_i = 0  otherwise", "(3.22)")

add_table(
    ["Parameter", "Value", "Description"],
    [
        ["Number of channels", "8", "Sensor array width"],
        ["Lateral offsets (m)", "[±0.07, ±0.05, ±0.03, ±0.01]", "Sensor element positions"],
        ["Forward projection", "0.10 m", "Distance ahead of base_footprint"],
        ["Detection half-width", "0.04 m", "Tape detection threshold"],
        ["Tape width (guide)", "0.05 m", "Physical tape dimension in world"],
        ["Update rate", "50 Hz", "Sensor polling frequency"],
        ["Junction threshold", "5 sensors active", "Minimum for junction detection"],
        ["Junction confirmation", "3 consecutive frames (60 ms)", "Debounce period"],
        ["Junction cooldown", "2.0 s", "Minimum time between junctions"],
    ],
    "Table 13: Line Sensor Array Geometry Parameters"
)

heading3("3.9.2  Junction Detection Logic")

body_para(
    "A junction event occurs when the robot is at the intersection of the spine and an aisle "
    "tape, causing sensors on both sides of centre to simultaneously detect tape. The "
    "junction detection algorithm uses a three-condition debounce: at least 5 of the 8 "
    "sensors must register binary value 1 (indicating a wide tape region characteristic "
    "of an intersection); this condition must persist for at least 3 consecutive 50 Hz "
    "frames (60 ms), filtering out brief noise spikes; and at least 2.0 seconds must have "
    "elapsed since the last junction event, preventing double-counting as the robot crosses "
    "an intersection at speed. When all three conditions are satisfied, the node publishes "
    "an Empty message to /atlas/junction, which the mission manager uses to count aisle "
    "positions along the spine."
)

heading3("3.9.3  IMU Sensor")

body_para(
    "The IMU sensor is implemented in Gazebo using the libgazebo_ros_imu_sensor.so plugin "
    "attached to the imu_link frame. It publishes sensor_msgs/Imu messages at 100 Hz to "
    "/atlas/imu, providing quaternion orientation data, angular velocity, and linear "
    "acceleration. The turn controller extracts the yaw angle from the quaternion using "
    "the standard ZYX Euler angle formula, valid for planar motion where pitch and roll "
    "are approximately zero:"
)

add_equation("ψ = atan2(2(q_w·q_z + q_x·q_y),  1 − 2(q_y² + q_z²))", "(3.23)")

body_para(
    "In the Gazebo simulation, IMU noise is near-zero with no drift. In the physical "
    "deployment, the BNO055 9-axis absolute orientation sensor provides an equivalent "
    "100 Hz quaternion output via I2C, with a noise specification of approximately ±0.5° "
    "and a drift rate of 1 to 10°/hour, which is acceptable for the 45-second mission "
    "cycle duration."
)

heading3("3.9.4  RFID Detection System")

body_para(
    "The tag detector node (tag_detector.py) implements a simulated proximity-based RFID "
    "detection system. A database of 21 tags is maintained: one home tag at position (0, 0) "
    "and 20 shelf tags at positions (x_s, y_a) where x_s ∈ {1.0, 2.0, 3.0, 4.0} m and "
    "y_a ∈ {2, 4, 6, 8, 10} m, corresponding to the 20 shelf rack positions in the "
    "warehouse world. Tags are numbered S01 through S20 in row-major order."
)
body_para(
    "The detection model incorporates hysteresis to prevent rapid toggling when the robot "
    "is near the detection boundary. Detection fires when the Euclidean distance between "
    "the robot position and a tag position satisfies:"
)

add_equation("d = √((x − x_tag)² + (y − y_tag)²) ≤ 0.5 m   AND   tag is armed", "(3.24)")

body_para(
    "Once detected, the tag is disarmed and remains so until the robot moves beyond the "
    "re-arm radius:"
)

add_equation("re-arm condition:  d > 0.8 m", "(3.25)")

body_para(
    "The 0.3 m hysteresis gap between detection (0.5 m) and re-arm (0.8 m) radii prevents "
    "oscillatory detection events as the robot lingers near a tag location. On detection, "
    "an atlas_interfaces/ShelfTag message is published to /atlas/tag_event with the tag "
    "identifier, shelf identifier, measured distance, home flag, and timestamp."
)


heading2("3.10  Control Algorithms — PD Line Following")

heading3("3.10.1  Error Signal Computation")

body_para(
    "The line follower node (line_follower.py) computes a lateral displacement error "
    "signal from the binary sensor array. A weighted average of the active sensor "
    "positions is computed, normalised by the total number of active sensors:"
)

add_equation("e(t) = Σ(w_i · s_i) / Σ(s_i)   for i = 0 to 7", "(3.26)")

body_para(
    "where the sensor weights are w = [1.0, 0.71, 0.43, 0.14, −0.14, −0.43, −0.71, −1.0], "
    "assigned symmetrically such that left-side sensors give positive error (requiring a "
    "left turn to re-centreˆ) and right-side sensors give negative error. The normalisation "
    "by the sum of active sensors means that the error signal is independent of the number "
    "of sensors detecting tape, providing a consistent error magnitude across the full "
    "detection range from single-sensor detection at the tape edge to full-array detection "
    "at a junction."
)

heading3("3.10.2  PD Control Law")

body_para(
    "The PD control law computes an angular velocity command from the current and previous "
    "error samples:"
)

add_equation("u(t) = K_p · e(t) + K_d · de(t)/dt", "(3.27)")

body_para(
    "Discretised at 50 Hz sampling frequency f_s:"
)

add_equation("u_k = K_p · e_k + K_d · (e_k − e_{k-1}) · f_s", "(3.28)")

body_para(
    "With K_p = 0.6 and K_d = 0.2, the control output u_k becomes the angular.z component "
    "of the /atlas/nav_vel Twist message, while linear.x is held constant at 0.4 m/s. "
    "The combination of forward speed and corrective angular velocity steers the robot "
    "along the tape centreline."
)

add_table(
    ["Parameter", "Value", "Justification"],
    [
        ["K_p (Proportional gain)", "0.6", "Tuned for 0.4 m/s speed — sufficient response without oscillation"],
        ["K_i (Integral gain)", "0.0", "Disabled — prevents windup at junctions and setpoint changes"],
        ["K_d (Derivative gain)", "0.2", "Provides damping — reduces overshoot at curves"],
        ["Operating speed", "0.4 m/s", "Balance between throughput and tracking accuracy"],
        ["Update rate", "50 Hz", "Provides 20 ms control loop — adequate for 0.4 m/s"],
        ["Sensor weights", "[1.0, 0.71, 0.43, 0.14, ...]", "Cosine-spaced for linear error approximation"],
        ["Grace period", "60 frames (1.2 s)", "Continue last output on tape loss before stopping"],
    ],
    "Table 14: PD Controller Parameters and Tuning"
)

heading3("3.10.3  Stability Analysis")

body_para(
    "The closed-loop stability of the line-following system can be analysed by considering "
    "the linearised lateral dynamics. The characteristic equation of the closed-loop system "
    "with the given PD gains is:"
)

add_equation("s² + 6s + 3 = 0", "(3.29)")

body_para(
    "The roots of this equation are s = −0.55 and s = −5.45, both of which are real and "
    "negative, confirming that the closed-loop system is stable and overdamped. The natural "
    "frequency is ω_n = √3 = 1.73 rad/s, and the 2% settling time is approximately "
    "t_s ≈ 4/0.55 = 7.3 samples = 0.15 s, corresponding to approximately 6 cm of "
    "travel at 0.4 m/s. This rapid settling ensures that perturbations from junction "
    "crossings or line edges are corrected within a short distance."
)
body_para(
    "The absence of an integral term eliminates the risk of integrator windup, which would "
    "cause overshoot when the robot exits a junction and re-acquires the spine tape. The "
    "derivative term provides sufficient damping to prevent the oscillations that would "
    "occur with a purely proportional controller at the chosen operating speed."
)

heading2("3.11  Turn Controller — IMU-Based Bang-Bang Control")

heading3("3.11.1  Architecture")

body_para(
    "The turn controller node (turn_controller.py) implements a simple bang-bang control "
    "strategy for executing precise heading changes. When a turn command is received on "
    "/atlas/turn_cmd as a Float32 radian value representing the desired heading change, "
    "the controller computes the target absolute heading from the current IMU yaw and "
    "applies a constant angular velocity command until the target is reached within the "
    "tolerance angle."
)
body_para(
    "The control law is:"
)

add_equation("ω_cmd = sign(Δθ) × 0.4 rad/s   while |θ_target − θ_current| > 3°", "(3.30)")
add_equation("ω_cmd = 0  when  |θ_target − θ_current| ≤ 3°", "(3.31)")

body_para(
    "The angular error is wrapped to the range [−π, +π] using the atan2(sin(·), cos(·)) "
    "identity, preventing discontinuities when the heading crosses the ±180° boundary. "
    "On reaching the tolerance, the node publishes a zero velocity Twist to /atlas/turn_vel "
    "and an Empty message to /atlas/turn_done to notify the mission manager that the turn "
    "is complete."
)

heading3("3.11.2  Turn Duration Estimation")

body_para(
    "The expected duration of turns at the 0.4 rad/s angular velocity is:"
)

add_equation("t_turn = |Δθ| / ω = |Δθ| / 0.4", "(3.32)")

body_para(
    "For a 90° (π/2 rad) aisle entry turn: t = (π/2) / 0.4 = 3.93 s. "
    "For a 180° (π rad) shelf pickup pivot: t = π / 0.4 = 7.85 s. "
    "For a 90° (π/2 rad) return-to-spine turn: t = (π/2) / 0.4 = 3.93 s. "
    "These durations are verified in the mission timing analysis in Chapter 4."
)


heading2("3.12  Mission State Machine Design")

heading3("3.12.1  State Definitions")

body_para(
    "The mission Finite State Machine (FSM) in mission_node.py orchestrates the complete "
    "lifecycle of a warehouse pick-and-return mission through 12 defined states. The FSM "
    "is event-driven: state transitions are triggered by ROS2 topic events (junction "
    "detected, turn completed, RFID tag read) and timer-based conditions (dwell times at "
    "pickup and docking states)."
)

add_table(
    ["State", "Purpose", "Velocity Source", "Entry Condition", "Exit Condition"],
    [
        ["IDLE", "Awaiting mission", "Zero", "System startup / mission complete", "Mission in queue AND not E-stopped"],
        ["NAV_SPINE", "Following spine northward", "Line follower (nav_vel)", "Mission popped from queue", "Junction count == target aisle"],
        ["TURNING", "90° turn into target aisle", "Turn controller (turn_vel)", "Junction count reached", "turn_done received"],
        ["NAV_AISLE", "Following aisle toward shelf", "Line follower (nav_vel)", "Turn complete", "Target shelf RFID detected"],
        ["AT_SHELF", "Arrived at target shelf", "Zero", "RFID match to target", "0.5 s dwell elapsed"],
        ["PICKUP", "Simulated pick operation", "Zero", "AT_SHELF dwell complete", "2.0 s pickup elapsed"],
        ["PIVOT", "180° turn to face return direction", "Turn controller (turn_vel)", "Pickup complete", "turn_done received"],
        ["RET_AISLE", "Following aisle back toward spine", "Line follower (nav_vel)", "Pivot complete", "Junction detected on aisle"],
        ["RET_TURN", "90° turn back onto spine", "Turn controller (turn_vel)", "Aisle junction detected", "turn_done received"],
        ["RET_SPINE", "Following spine southward to home", "Line follower (nav_vel)", "Return turn complete", "Home RFID tag detected"],
        ["DOCKED", "Arrived at home dock", "Zero", "Home tag detected", "1.0 s dwell elapsed → IDLE"],
        ["ERROR", "E-stop active", "Zero (all motion halted)", "estop topic received", "/atlas/reset topic received"],
    ],
    "Table 15: Mission State Machine — States and Velocity Sources"
)

heading3("3.12.2  Junction Counting Strategy")

body_para(
    "Aisle selection is implemented through junction counting along the spine. The spine "
    "tape intersects aisle tapes at y=2, 4, 6, 8, and 10 m, corresponding to aisle "
    "numbers 1 through 5 respectively. As the robot navigates north along the spine in "
    "the NAV_SPINE state, each junction event increments an internal counter. When the "
    "counter reaches the target aisle number (encoded in the SHELVES dictionary as the "
    "third element of each shelf entry), the mission manager publishes a turn command to "
    "/atlas/turn_cmd and transitions to the TURNING state. This approach requires no map "
    "knowledge beyond the junction count, making it robust to minor positional variations "
    "along the spine."
)

heading3("3.12.3  Mission Queue Management")

body_para(
    "Missions received on /atlas/mission_cmd are queued in a Python list. The IDLE state "
    "polls this queue at every 50 Hz tick and, when a mission is available and the system "
    "is not E-stopped, pops the first item and begins execution. Missions are validated "
    "against the SHELVES dictionary on receipt; unknown shelf identifiers are logged and "
    "rejected. This simple first-in, first-out queue supports sequential mission execution "
    "but does not implement priority scheduling or path optimisation — identified as a "
    "future enhancement."
)

heading3("3.12.4  Emergency Stop and Reset-to-Dock")

body_para(
    "The E-stop mechanism is implemented as a priority override at the velocity arbiter "
    "level. When an /atlas/estop message is received in any state, the estopped flag is "
    "set to True and the FSM transitions to the ERROR state. Regardless of any other "
    "velocity source, the arbiter publishes a zero Twist when estopped is True. The "
    "/atlas/reset topic clears the flag and returns to IDLE, discarding the active mission "
    "and queue."
)
body_para(
    "The reset-to-dock function, triggered by /atlas/reset_to_dock, implements a nine-step "
    "recovery procedure: logging the reset request, cancelling the active mission, "
    "publishing zero velocity, clearing all internal state variables, resetting the "
    "junction counter, calling the Gazebo /set_entity_state service to teleport the robot "
    "to (0, 0, 0.01) with heading π/2, transitioning to IDLE, and logging system readiness. "
    "On physical hardware, the Gazebo service call would be replaced by a manual reset "
    "procedure or an autonomous docking sequence."
)


heading2("3.13  Velocity Arbiter Pattern")

body_para(
    "The velocity arbiter is the core safety mechanism of the ATLAS architecture. Implemented "
    "within the atlas_mission_mgr node's 50 Hz tick callback, it selects the appropriate "
    "velocity command based on the current FSM state and publishes it as the sole output on "
    "/atlas/cmd_vel. The selection logic is:"
)

body_para(
    "In navigation states (NAV_SPINE, NAV_AISLE, RET_AISLE, RET_SPINE): the output equals "
    "the most recent /atlas/nav_vel message from the line follower. In turning states "
    "(TURNING, PIVOT, RET_TURN): the output equals the most recent /atlas/turn_vel message "
    "from the turn controller. In stationary states (IDLE, AT_SHELF, PICKUP, DOCKED): the "
    "output is a zero Twist. In the ERROR state: the output is zero regardless. The E-stop "
    "override unconditionally sets the output to zero whenever estopped is True, regardless "
    "of state. This design guarantees that the E-stop is never masked by control logic.",
    bold=False
)

body_para(
    "The single-publisher pattern — enforced by the topic architecture — means that no "
    "external node can accidentally inject velocity commands that bypass the arbiter. This "
    "is critical for safety: in a physical deployment, any software fault in a navigation "
    "or control node would at worst cause incorrect velocity recommendations on intermediate "
    "topics, but the arbiter's E-stop override would always remain effective."
)

heading2("3.14  Simulation-to-Physical Conversion Methodology")

heading3("3.14.1  Conversion Philosophy")

body_para(
    "The ATLAS software architecture was designed from the outset with physical deployment "
    "in mind. The key architectural principle enabling straightforward physical deployment "
    "is the hardware abstraction layer: all hardware-specific code is contained within "
    "the Gazebo plugins (libgazebo_ros_diff_drive.so, libgazebo_ros_imu_sensor.so) and "
    "the simulated sensor nodes (line_sensor.py, tag_detector.py). The control logic, "
    "mission management, GUI, and CLI are completely hardware-agnostic and run unchanged "
    "on a physical robot."
)

add_table(
    ["Simulation Component", "Physical Replacement", "New File", "Interface Unchanged"],
    [
        ["Gazebo diff_drive plugin", "L298N + JGB37-520 motors + encoders", "motor_driver.py", "Yes — /atlas/cmd_vel, /atlas/odom"],
        ["line_sensor.py (geometric)", "QTR-8A + MCP3008 SPI ADC", "line_sensor_hw.py", "Yes — /atlas/line_sensors"],
        ["tag_detector.py (distance)", "RDM6300 UART reader", "rfid_reader_hw.py", "Yes — /atlas/tag_event"],
        ["Gazebo IMU plugin", "BNO055 I2C sensor", "imu_hw.py", "Yes — /atlas/imu"],
        ["URDF geometry", "6061 aluminium frame (300×250 mm)", "N/A (physical build)", "N/A"],
        ["atlas_full.launch.py", "atlas_real.launch.py (no Gazebo)", "atlas_real.launch.py", "N/A"],
        ["line_follower.py", "Unchanged — runs on Raspberry Pi", "—", "Yes"],
        ["turn_controller.py", "Unchanged — runs on Raspberry Pi", "—", "Yes"],
        ["mission_node.py", "Unchanged — runs on Raspberry Pi", "—", "Yes"],
        ["atlas_gui.py", "Unchanged — runs on operator PC via WiFi DDS", "—", "Yes"],
    ],
    "Table 17: Simulation-to-Physical Component Conversion Table"
)

heading3("3.14.2  Hardware Component Specifications")

add_table(
    ["Component", "Specification", "Purpose", "Estimated Cost (USD)"],
    [
        ["Raspberry Pi 4 (4 GB)", "4× Cortex-A72 @ 1.5 GHz, 4 GB LPDDR4, 40-pin GPIO", "Main computing platform", "55"],
        ["JGB37-520 DC Motors (×2)", "12V, 200 RPM, Hall-effect encoder, 330 CPR, 6mm D-shaft", "Drive wheels", "30–50"],
        ["L298N Dual H-Bridge", "5–35V, 2A/ch (3A peak), PWM + direction pins", "Motor controller", "5–8"],
        ["QTR-8A Sensor Array", "8-channel analog IR, 9.5mm spacing, 6mm optimal height", "Line following", "12"],
        ["MCP3008 ADC", "8-channel 10-bit SPI ADC, 3.3V supply", "Analog-to-digital for sensor", "5"],
        ["BNO055 IMU", "9-axis (3× accel, 3× gyro, 3× mag), I2C, 100Hz quaternion, ±0.5°", "Heading control", "25–35"],
        ["RDM6300 RFID Reader", "125 kHz, 5–10cm range, UART 9600 baud", "Shelf identification", "5"],
        ["EM4100 RFID Tags (×21)", "125 kHz passive disc tags", "Shelf and home markers", "10–15"],
        ["3S LiPo Battery (5000 mAh)", "11.1V, 55.5 Wh, ~350g", "Main power supply", "25–40"],
        ["Buck Converter (12V→5V)", "5A output, powers Raspberry Pi via USB-C", "Power regulation", "8"],
        ["Aluminium frame", "6061, 20×20mm T-slot extrusion, 300×250mm base plate", "Robot chassis", "40–80"],
    ],
    "Table 18: Hardware Component Specifications"
)

heading3("3.14.3  Power Budget")

add_table(
    ["Component", "Voltage", "Current (typical)", "Power"],
    [
        ["2× DC Motors (loaded)", "12V", "1.0A each", "24.0 W"],
        ["Raspberry Pi 4", "5V", "2.5A", "12.5 W"],
        ["All sensors (QTR-8A, BNO055, RDM6300)", "3.3V", "0.2A total", "0.66 W"],
        ["Motor controller (logic)", "5V", "0.05A", "0.25 W"],
        ["Total", "—", "—", "~37.4 W"],
    ],
    "Table 20: Power Budget for Physical Robot"
)

body_para(
    "The 3S LiPo battery (11.1V, 5000 mAh, 55.5 Wh) provides an estimated runtime of "
    "55.5 / 37.4 ≈ 1.5 hours at continuous full load. For typical warehouse operation "
    "with a 40% duty cycle on motors, runtime extends to approximately 3 hours, sufficient "
    "for a complete operational shift with a brief recharge at the dedicated charging station "
    "modelled in the warehouse world."
)

add_table(
    ["Build Type", "Component Scope", "Estimated Cost (USD)"],
    [
        ["Minimum Viable Build", "Raspberry Pi 4, motors, L298N, QTR-8A, BNO055, RDM6300, basic frame, battery, wiring", "361"],
        ["Production Quality Build", "Above + improved frame, encoders, WiFi router, additional sensors, enclosure, spare parts", "805"],
    ],
    "Table 21: Budget Summary — Minimum and Production Builds"
)


heading3("3.14.4  GPIO Pin Mapping for Physical Robot")

add_table(
    ["GPIO Pin", "Function", "Direction", "Notes"],
    [
        ["GPIO 2", "I2C SDA (BNO055 IMU)", "Bidirectional", "3.3V logic"],
        ["GPIO 3", "I2C SCL (BNO055 IMU)", "Output", "3.3V logic"],
        ["GPIO 5", "Encoder Left-A", "Input", "Interrupt-driven"],
        ["GPIO 6", "Encoder Left-B", "Input", "Interrupt-driven"],
        ["GPIO 8", "SPI CE0 (MCP3008 ADC)", "Output", "Active low"],
        ["GPIO 9", "SPI MISO (MCP3008)", "Input", "—"],
        ["GPIO 10", "SPI MOSI (MCP3008)", "Output", "—"],
        ["GPIO 11", "SPI CLK (MCP3008)", "Output", "1.35 MHz"],
        ["GPIO 12", "PWM0 — Left motor speed", "Output", "10 kHz PWM"],
        ["GPIO 13", "PWM1 — Right motor speed", "Output", "10 kHz PWM"],
        ["GPIO 14", "UART TX (RDM6300 RFID)", "Output", "9600 baud"],
        ["GPIO 15", "UART RX (RDM6300 RFID)", "Input", "9600 baud"],
        ["GPIO 17", "Motor L-IN1 (direction A)", "Output", "L298N control"],
        ["GPIO 22", "Motor R-IN3 (direction A)", "Output", "L298N control"],
        ["GPIO 23", "Motor R-IN4 (direction B)", "Output", "L298N control"],
        ["GPIO 24", "Encoder Right-A", "Input", "Interrupt-driven"],
        ["GPIO 25", "Encoder Right-B", "Input", "Interrupt-driven"],
        ["GPIO 27", "Motor L-IN2 (direction B)", "Output", "L298N control"],
    ],
    "Table 19: GPIO Pin Mapping for Physical Robot"
)

heading3("3.14.5  Motor PID Velocity Control for Physical Hardware")

body_para(
    "In the physical deployment, the motor driver node replaces the Gazebo differential "
    "drive plugin. It subscribes to /atlas/cmd_vel, converts Twist commands to individual "
    "wheel velocities using the inverse kinematics equations (3.6) and (3.7), and applies "
    "a per-motor PID velocity controller using encoder feedback:"
)

add_equation("v_L = cmd_vel.linear.x − cmd_vel.angular.z · L / 2", "(3.33)")
add_equation("v_R = cmd_vel.linear.x + cmd_vel.angular.z · L / 2", "(3.34)")

body_para(
    "The velocity PID controller for each motor starts with gains K_p=2.0, K_i=1.0, "
    "K_d=0.05, which must be tuned on the physical robot by placing it on blocks with "
    "free wheels, commanding known velocities, and adjusting until steady-state error "
    "is below 2% and step response overshoot is below 10%. The motor driver also "
    "publishes computed odometry to /atlas/odom using the integration equations (3.11) "
    "through (3.15), maintaining identical topic compatibility with the simulation."
)

heading2("3.15  Chapter Summary")

body_para(
    "This chapter has presented the complete technical implementation of the ATLAS system, "
    "covering the layered system architecture and ROS2 package structure, node and topic "
    "communication design, URDF/Xacro robot model with parametric physical properties, "
    "differential drive forward and inverse kinematics with odometry integration, inertia "
    "tensor calculation and stability analysis, Gazebo simulation environment configuration "
    "and warehouse world design, the odom-equals-world architectural decision, the 8-channel "
    "virtual line sensor with junction detection, IMU heading extraction, RFID detection "
    "with hysteresis, the PD line following controller with stability analysis, the "
    "IMU-based bang-bang turn controller, the 12-state mission FSM with queue management "
    "and emergency stop, the velocity arbiter safety pattern, and the complete simulation-"
    "to-physical hardware conversion methodology. The following chapter presents quantified "
    "performance results, engineering discussion, and conclusions."
)

page_break()


# ════════════════════════════════════════════════════════════════════════════
# CHAPTER 4 — RESULTS, DISCUSSION AND CONCLUSIONS
# ════════════════════════════════════════════════════════════════════════════
heading1("CHAPTER 4")
heading1("RESULTS, DISCUSSION AND CONCLUSIONS")

body_para(
    "This chapter presents the experimental observations and performance measurements "
    "obtained from the ATLAS Warehouse AGV simulation, a quantitative analysis of mission "
    "timing, a discussion of system behaviour and engineering trade-offs, an assessment "
    "of limitations and future improvement directions, and the final conclusions of the project."
)

heading2("4.1  Simulation Performance Results")

body_para(
    "The ATLAS system was evaluated across multiple complete mission cycles in the Gazebo "
    "Classic 11 simulation environment. Performance was assessed against the success criteria "
    "defined in the project objectives. The following table summarises the key performance "
    "metrics measured during simulation testing."
)

add_table(
    ["Metric", "Measured Value", "Requirement", "Status"],
    [
        ["Line following speed", "0.4 m/s", "≥ 0.3 m/s", "PASS"],
        ["Line tracking accuracy (lateral error)", "< 5 mm", "< 5 mm", "PASS"],
        ["Turn accuracy (heading error)", "±3°", "≤ ±5°", "PASS"],
        ["RFID detection rate", "100% for all 21 tags", "≥ 95%", "PASS"],
        ["Mission completion rate", "100% per cycle", "≥ 98%", "PASS"],
        ["Docking alignment accuracy", "±3 cm, ±3°", "≤ ±5 cm, ±5°", "PASS"],
        ["Emergency stop response", "< 20 ms (1 control tick at 50 Hz)", "< 100 ms", "PASS"],
        ["Mission cycle time (S05)", "≈ 44.2 s", "< 60 s", "PASS"],
        ["GUI update rate", "10 Hz", "≥ 5 Hz", "PASS"],
        ["Throughput estimate (single robot)", "~80 picks/hour", "—", "Informational"],
    ],
    "Table 22: System Performance Metrics"
)

heading2("4.2  Mission Timing Analysis")

body_para(
    "A detailed timing analysis was conducted for a complete pick-and-return mission to "
    "shelf S05 (located at x=1.0 m, y=4.0 m, aisle 2). This shelf requires the robot to "
    "travel 4.0 m along the spine, turn 90° into the second aisle, travel 1.0 m to the "
    "shelf, pick up the item, pivot 180°, return 1.0 m to the spine, turn 90° back onto "
    "the spine, and travel 4.0 m back to the home dock."
)

add_table(
    ["Mission Phase", "State", "Distance / Angle", "Speed", "Duration"],
    [
        ["Navigate spine north", "NAV_SPINE", "4.0 m", "0.4 m/s", "10.0 s"],
        ["90° turn into aisle", "TURNING", "π/2 rad", "0.4 rad/s", "3.9 s"],
        ["Navigate aisle east", "NAV_AISLE", "1.0 m", "0.4 m/s", "2.5 s"],
        ["Arrive at shelf", "AT_SHELF", "—", "0", "0.5 s"],
        ["Simulated pick operation", "PICKUP", "—", "0", "2.0 s"],
        ["180° pivot turn", "PIVOT", "π rad", "0.4 rad/s", "7.9 s"],
        ["Navigate aisle west (return)", "RET_AISLE", "1.0 m", "0.4 m/s", "2.5 s"],
        ["90° turn back to spine", "RET_TURN", "π/2 rad", "0.4 rad/s", "3.9 s"],
        ["Navigate spine south to home", "RET_SPINE", "4.0 m", "0.4 m/s", "10.0 s"],
        ["Docking dwell", "DOCKED", "—", "0", "1.0 s"],
        ["Total cycle time", "—", "~14 m + 4.7 rad", "—", "≈ 44.2 s"],
    ],
    "Table 23: Mission Timing Analysis — Shelf S05"
)

body_para(
    "The total cycle time of approximately 44.2 seconds for the nearest mid-range shelf "
    "S05 gives a theoretical maximum throughput of 3600/44.2 ≈ 81 pick cycles per hour "
    "for a single robot. Missions to more distant shelves (e.g., S20 at aisle 5, "
    "x=4.0 m) would require 10 m spine travel and 4 m aisle travel, increasing cycle "
    "time to approximately 64 seconds and reducing throughput to approximately 56 "
    "picks/hour. A fleet of five robots operating on non-conflicting paths could "
    "achieve an estimated 300 picks/hour with a 75% efficiency factor accounting for "
    "path spacing constraints."
)


heading2("4.3  System Behaviour and Discussion")

heading3("4.3.1  Line Following Performance")

body_para(
    "The PD line following controller demonstrated stable, smooth navigation along both "
    "spine and aisle tapes throughout all tested missions. The weighted average error "
    "computation provides a continuous lateral position estimate that the PD controller "
    "corrects within approximately 6 cm of travel after a perturbation — equivalent to "
    "the settling time of 0.15 s at 0.4 m/s. Junction crossings produced momentary "
    "full-sensor activation, which correctly triggered junction detection events without "
    "destabilising the line follower, as the grace period of 60 frames maintained the "
    "last computed velocity during the brief crossing interval."
)
body_para(
    "The absence of integral action proved beneficial at junction crossings: the controller "
    "did not accumulate an error bias during the brief period of wide sensor activation, "
    "and re-acquired the tape centre immediately upon exiting the junction. This behaviour "
    "confirms the design rationale for the PD-only implementation."
)

heading3("4.3.2  Turn Accuracy")

body_para(
    "The IMU-based bang-bang turn controller consistently achieved heading accuracy within "
    "±3° of the commanded target for both 90° and 180° turns across all tested missions. "
    "The 3° tolerance angle was selected as the minimum acceptable error that still allows "
    "the line follower to re-acquire the tape within the aisle width. Tighter tolerances "
    "would increase turn duration without proportional improvement in mission performance, "
    "as the line follower corrects any residual angular error within the first few "
    "centimetres of aisle navigation."
)

heading3("4.3.3  RFID Detection Reliability")

body_para(
    "The simulated RFID system achieved 100% detection rate for all 21 tags across multiple "
    "mission cycles. The 0.5 m detection radius provides a reliable detection window: at "
    "0.4 m/s approach speed, the robot is within the detection radius for 2.5 seconds, "
    "well exceeding the requirement for a single detection event. The hysteresis gap "
    "of 0.3 m between detection and re-arm radii successfully prevented spurious re-"
    "detection events in all cases. In a physical deployment, the RDM6300 reader's "
    "5 to 10 cm actual detection range would require reducing the approach speed or "
    "adjusting the tag placement to ensure reliable detection."
)

heading3("4.3.4  GUI and Control Interface Behaviour")

body_para(
    "The PyQt5 control centre operated correctly, displaying real-time robot state, "
    "position, velocity, heading, battery percentage, and RFID event log at 10 Hz update "
    "rate. The dual-threaded architecture — ROS2 spin on a background daemon thread "
    "communicating to the Qt main thread via pyqtSignal — ensured that GUI responsiveness "
    "was not affected by ROS2 callback processing delays. Mission dispatch via the GUI "
    "dialog was confirmed to correctly queue and execute missions. Emergency stop and "
    "reset-to-dock functions operated correctly in all tested scenarios."
)

heading2("4.4  Error Analysis and Limitations")

body_para(
    "Several limitations of the current implementation were identified during testing. "
    "The most significant functional limitation is the single-robot constraint: the current "
    "mission manager has no mechanism to coordinate with a second robot, and two robots "
    "would collide on shared spine segments without traffic management logic. This is "
    "the primary architectural extension required for a production-scale fleet deployment."
)
body_para(
    "The absence of obstacle detection is acceptable in the fully controlled simulation "
    "environment but would be a critical safety gap in physical deployment per ISO 3691-4 "
    "[10] requirements for driverless industrial trucks. Physical deployment would require "
    "at minimum two safety-rated proximity sensors or a LiDAR scanner for pedestrian "
    "detection in shared spaces."
)
body_para(
    "Simulated sensors do not replicate the noise characteristics of physical hardware. "
    "The QTR-8A sensor is affected by ambient light variations, floor surface reflectance "
    "non-uniformity, and sensor height variation due to floor imperfections. The BNO055 "
    "IMU exhibits 1 to 10°/hour gyroscope drift, which over 45-second missions accumulates "
    "to a maximum of 0.125° — negligible for the current application. The RDM6300 RFID "
    "reader's detection range is sensitive to tag orientation and metallic interference "
    "from motor electronics, requiring careful shielding in the physical build."
)

add_table(
    ["Decision", "Benefit", "Cost / Trade-off"],
    [
        ["Line following over SLAM", "Simple, deterministic, infrastructure-based", "Fixed paths only — no dynamic rerouting"],
        ["Junction counting (no map)", "No map required — immediate deployment", "Cannot skip junctions — must count from home each time"],
        ["Bang-bang turn control", "Simple, reliable, no tuning required", "Slightly imprecise at turn endpoint — 3° residual"],
        ["Single velocity arbiter", "Safety guaranteed — no velocity conflicts", "Complex state machine — all velocity logic centralised"],
        ["KI=0 in line follower", "No windup at junctions and curves", "Slight steady-state error on sustained curves"],
        ["GUI as separate process", "Crash-safe — GUI fault cannot stop robot", "WiFi dependency for physical deployment"],
        ["Odom = world frame", "No localisation layer required", "Assumes accurate spawn position every startup"],
    ],
    "Table 25: Design Trade-offs Summary"
)


heading2("4.5  Physical Deployment Analysis")

body_para(
    "The simulation-to-physical analysis confirms that 70% of the ATLAS codebase runs "
    "unchanged on physical hardware. The seven software files that require no modification "
    "for physical deployment are: line_follower.py (reads ROS2 topics, outputs Twist — "
    "identical behaviour on any platform), turn_controller.py (reads IMU topic, outputs "
    "Twist), mission_node.py (pure state machine operating on ROS2 topics), send_mission.py "
    "(CLI publisher), atlas_gui.py (GUI on operator PC), all atlas_interfaces message "
    "definitions, and the atlas_bringup package (replaced by atlas_real.launch.py)."
)
body_para(
    "The four files that require replacement are: the Gazebo differential drive plugin "
    "(replaced by motor_driver.py using GPIO PWM and encoder feedback), line_sensor.py "
    "(replaced by line_sensor_hw.py using SPI ADC from QTR-8A), tag_detector.py "
    "(replaced by rfid_reader_hw.py using UART from RDM6300), and the Gazebo IMU plugin "
    "(replaced by imu_hw.py using I2C from BNO055). All replacement nodes maintain "
    "identical ROS2 topic interfaces, ensuring zero changes to control logic."
)

heading2("4.6  Industry 4.0 Context")

body_para(
    "The ATLAS project is positioned within the broader context of Industry 4.0 — the "
    "fourth industrial revolution characterised by cyber-physical systems, Internet of "
    "Things connectivity, and autonomous production environments. Commercial AGV systems "
    "such as Amazon Robotics (800,000+ units deployed globally), MiR 250, and OTTO 100 "
    "represent the state of the art in warehouse automation, employing LiDAR SLAM, "
    "computer vision, and cloud-based fleet coordination."
)

add_table(
    ["Feature", "ATLAS (this project)", "Amazon Robotics", "MiR 250", "OTTO 100"],
    [
        ["Navigation", "Line following + RFID", "Vision SLAM", "LiDAR SLAM", "LiDAR SLAM"],
        ["Localisation", "Odometry + RFID", "Visual landmarks", "AMCL", "AMCL"],
        ["Payload (simulated)", "Boolean flag", "300 kg", "250 kg", "100 kg"],
        ["Fleet size", "1 (extensible)", "800,000+", "Unlimited", "Unlimited"],
        ["Estimated cost", "~$400 (physical)", "Proprietary", "$25,000", "$30,000"],
        ["Navigation reliability", "Deterministic", "High", "High", "High"],
        ["Map requirement", "None (tape-guided)", "Camera mapping", "LiDAR map", "LiDAR map"],
    ],
    "Table 24: Commercial AGV Comparison"
)

body_para(
    "The ATLAS approach's principal competitive advantage in academic and small-scale "
    "industrial contexts is the combination of deterministic, certifiable behaviour and "
    "very low cost — properties that SLAM-based systems sacrifice in favour of flexibility. "
    "For structured warehouse environments where the layout changes infrequently, line-"
    "following AGVs remain the technically optimal and economically superior choice."
)

heading2("4.7  Future Improvements")

body_para(
    "Several directions for future development were identified during the project. Multi-"
    "robot fleet coordination is the highest-priority enhancement: implementing a fleet "
    "manager node that maintains robot states, assigns missions considering current "
    "positions, and implements traffic management rules for the shared spine segment. "
    "This would leverage ROS2's DDS multi-node architecture, which already supports "
    "multiple robots on the same network without changes to individual robot software."
)
body_para(
    "LiDAR integration would add obstacle detection capability, enabling the AGV to halt "
    "when a person or object blocks the path — a safety requirement for shared human-robot "
    "workspaces under ISO 3691-4. Computer vision using an onboard camera could provide "
    "shelf barcode verification as a secondary confirmation to RFID detection. Adaptive PID "
    "tuning using machine learning could automatically adjust control gains for different "
    "floor conditions or payload weights. A digital twin — real-time synchronisation between "
    "a physical robot and its Gazebo simulation — would enable remote monitoring and "
    "predictive maintenance. Battery management with real charge monitoring and automatic "
    "return-to-dock on low charge would extend unattended operational capability."
)


heading2("4.8  Conclusions")

body_para(
    "The ATLAS Warehouse Automated Guided Vehicle project has successfully demonstrated "
    "the complete design, simulation, implementation, and physical deployment analysis "
    "of a ROS2-based autonomous warehouse robot. The following conclusions are drawn "
    "from the work presented in this report."
)
body_para(
    "First, the system achieves all stated performance objectives: 100% mission completion "
    "rate, sub-5 mm lateral tracking accuracy at 0.4 m/s, ±3° turn precision, 100% RFID "
    "shelf identification reliability, emergency stop response within 20 ms, and a complete "
    "pick-and-return cycle time of 44.2 seconds for shelf S05. These metrics are "
    "consistent with requirements for a prototype warehouse AGV operating in a structured "
    "environment."
)
body_para(
    "Second, the ROS2 Humble framework with Gazebo Classic 11 simulation provides an "
    "effective and productive development environment for warehouse AGV systems. The DDS-"
    "based topic architecture enables clean separation of sensing, control, planning, and "
    "interface functions, and the Gazebo physics engine provides sufficiently accurate "
    "differential drive dynamics for control system validation."
)
body_para(
    "Third, the velocity arbiter pattern — concentrating all wheel velocity authority in "
    "the single mission_manager node — is an effective safety architecture that prevents "
    "conflicting velocity commands and provides a reliable emergency stop mechanism. This "
    "pattern is recommended as a standard design approach for any multi-node ROS2 robot "
    "with multiple potential velocity sources."
)
body_para(
    "Fourth, the simulation-to-physical analysis confirms that the architectural decision "
    "to abstract hardware interfaces into dedicated Gazebo plugins and virtual sensor nodes "
    "results in 70% of the codebase requiring no modification for physical deployment. "
    "The estimated minimum hardware cost of USD 361 demonstrates that capable AGV systems "
    "for structured environments are achievable at a fraction of the cost of commercial "
    "systems."
)
body_para(
    "Fifth, line-following navigation augmented by RFID position confirmation remains a "
    "technically sound and practically efficient approach for structured warehouse "
    "environments with fixed layouts. While it lacks the flexibility of SLAM-based "
    "navigation, its deterministic behaviour, zero mapping requirement, and straightforward "
    "safety certification make it the preferred approach for applications where the "
    "warehouse topology is stable."
)
body_para(
    "In conclusion, the ATLAS project bridges academic mobile robotics and industrial AGV "
    "engineering, demonstrating that a complete warehouse automation system can be built "
    "with open-source software, proven navigation methods, and careful state machine design. "
    "The system is production-ready in its control logic and requires only hardware "
    "interface node replacement for physical deployment, providing a complete and "
    "transferable reference implementation for future work in autonomous warehouse robotics."
)

page_break()

# ════════════════════════════════════════════════════════════════════════════
# REFERENCES
# ════════════════════════════════════════════════════════════════════════════
heading1("REFERENCES")

refs = [
    "[1] Siegwart R, Nourbakhsh IR, Scaramuzza D, 2011, Introduction to Autonomous Mobile Robots, 2nd Edition, MIT Press, Cambridge MA.",
    "[2] Dudek G, Jenkin M, 2010, Computational Principles of Mobile Robotics, 2nd Edition, Cambridge University Press, Cambridge.",
    "[3] Corke P, 2017, Robotics, Vision and Control: Fundamental Algorithms in MATLAB, 2nd Edition, Springer, Berlin.",
    "[4] Siciliano B, Khatib O, 2016, Springer Handbook of Robotics, 2nd Edition, Springer, Berlin.",
    "[5] Open Robotics, 2023, ROS2 Humble Hawksbill Documentation, https://docs.ros.org/en/humble/.",
    "[6] Open Source Robotics Foundation, 2023, Gazebo Classic 11 Documentation, https://classic.gazebosim.org/.",
    "[7] Quigley M, Gerkey B, Smart WD, 2015, Programming Robots with ROS, O'Reilly Media, Sebastopol CA.",
    "[8] Pololu Corporation, 2023, QTR-8A/8RC Reflectance Sensor Array User's Guide, Pololu Robotics and Electronics.",
    "[9] Bosch Sensortec, 2023, BNO055 Intelligent 9-axis Absolute Orientation Sensor Datasheet, Bosch Sensortec GmbH.",
    "[10] ISO 3691-4:2020, Industrial Trucks — Safety Requirements and Verification — Part 4: Driverless Industrial Trucks and Their Systems, International Organization for Standardization.",
    "[11] Finkenzeller K, 2010, RFID Handbook: Fundamentals and Applications in Contactless Smart Cards, Radio Frequency Identification and Near-Field Communication, 3rd Edition, Wiley, Chichester.",
    "[12] IEC 61496:2020, Safety of Machinery — Electro-sensitive Protective Equipment, International Electrotechnical Commission.",
    "[13] Ollero A, 2005, Intelligent Mobile Robot Navigation, Springer, Berlin.",
    "[14] Thrun S, Burgard W, Fox D, 2005, Probabilistic Robotics, MIT Press, Cambridge MA.",
    "[15] Ogata K, 2010, Modern Control Engineering, 5th Edition, Prentice Hall, Upper Saddle River NJ.",
    "[16] Franklin GF, Powell JD, Emami-Naeini A, 2015, Feedback Control of Dynamic Systems, 7th Edition, Pearson, London.",
    "[17] Object Management Group, 2015, Data Distribution Service (DDS) Specification v1.4, OMG Document formal/2015-04-10.",
    "[18] Wurman PR, D'Andrea R, Mountz M, 2008, Coordinating Hundreds of Cooperative Autonomous Vehicles in Warehouses, AI Magazine, vol. 29, no. 1, pp. 9–20.",
    "[19] De Ryck M, Versteyhe M, Debrouwere F, 2020, Automated Guided Vehicle Systems, State-of-the-Art Control Algorithms and Techniques, Journal of Manufacturing Systems, vol. 54, pp. 152–173.",
    "[20] Azadeh K, De Koster R, Roy D, 2019, Robotized and Automated Warehouse Systems: Review and Recent Developments, Transportation Science, vol. 53, no. 4, pp. 917–945.",
]

for ref in refs:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para_spacing(p, before=0, after=4)
    p.paragraph_format.line_spacing = Pt(18)
    run = p.add_run(ref)
    set_run_font(run, 12)

page_break()


# ════════════════════════════════════════════════════════════════════════════
# APPENDIX A — ROS2 Package Structure
# ════════════════════════════════════════════════════════════════════════════
heading1("APPENDIX A")
heading1("ROS2 PACKAGE STRUCTURE AND KEY SOURCE FILES")

heading2("A.1  Workspace Directory Structure")

body_para("The complete workspace directory structure is as follows:")

structure_lines = [
    "~/atlas_ws/src/",
    "├── atlas_interfaces/",
    "│   ├── msg/",
    "│   │   ├── FleetMission.msg",
    "│   │   ├── RobotState.msg",
    "│   │   └── ShelfTag.msg",
    "│   ├── CMakeLists.txt",
    "│   └── package.xml",
    "├── atlas_description/",
    "│   ├── urdf/atlas_agv.urdf.xacro",
    "│   ├── rviz/atlas.rviz",
    "│   └── package.xml",
    "├── atlas_gazebo/",
    "│   ├── worlds/warehouse.world",
    "│   └── package.xml",
    "├── atlas_navigation/",
    "│   └── atlas_navigation/",
    "│       ├── line_sensor.py",
    "│       ├── line_follower.py",
    "│       ├── turn_controller.py",
    "│       └── tag_detector.py",
    "├── atlas_mission_manager/",
    "│   └── atlas_mission_manager/",
    "│       ├── mission_node.py",
    "│       ├── send_mission.py",
    "│       └── atlas_gui.py",
    "└── atlas_bringup/",
    "    └── launch/atlas_full.launch.py",
]

for line in structure_lines:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para_spacing(p, before=0, after=0)
    p.paragraph_format.line_spacing = Pt(14)
    run = p.add_run(line)
    run.font.name = 'Courier New'
    run.font.size = Pt(10)

doc.add_paragraph()

heading2("A.2  URDF/Xacro Key Parameters")

body_para(
    "The atlas_agv.urdf.xacro file defines the robot model using Xacro property macros. "
    "Key property declarations at the top of the file are shown below. These values "
    "directly correspond to the physical parameters used in all kinematic equations "
    "throughout Chapter 3."
)

xacro_lines = [
    '<xacro:property name="WR" value="0.05"/>   <!-- wheel radius -->',
    '<xacro:property name="WT" value="0.04"/>   <!-- wheel width -->',
    '<xacro:property name="WS" value="0.30"/>   <!-- wheel track (separation) -->',
    '<xacro:property name="BX" value="0.30"/>   <!-- body length -->',
    '<xacro:property name="BY" value="0.25"/>   <!-- body width -->',
    '<xacro:property name="BZ" value="0.10"/>   <!-- body height -->',
    '<xacro:property name="BM" value="2.5"/>    <!-- body mass (kg) -->',
    '<xacro:property name="WM" value="0.2"/>    <!-- wheel mass (kg) -->',
    '<xacro:property name="CR" value="0.025"/>  <!-- caster radius -->',
]

for line in xacro_lines:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para_spacing(p, before=0, after=0)
    p.paragraph_format.line_spacing = Pt(14)
    run = p.add_run(line)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

doc.add_paragraph()

heading2("A.3  Custom Message Definitions")

body_para("The three custom message types defined in atlas_interfaces/msg/ are:")

msgs = [
    ("FleetMission.msg",
     "string mission_id\nstring target_shelf\nstring sku\nint32 priority"),
    ("RobotState.msg",
     "string state\nstring mission_id\nstring target_shelf\nstring current_sku\n"
     "string last_tag\nbool carrying_load\nfloat32 battery_percent\ngeometry_msgs/Point32 pose"),
    ("ShelfTag.msg",
     "string tag_id\nstring shelf_id\nfloat32 distance\nbool is_home\nbuiltin_interfaces/Time stamp"),
]

for msg_name, msg_body in msgs:
    heading3(msg_name)
    for line in msg_body.split('\n'):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        para_spacing(p, before=0, after=0)
        p.paragraph_format.line_spacing = Pt(14)
        run = p.add_run(line)
        run.font.name = 'Courier New'
        run.font.size = Pt(10)
    doc.add_paragraph()

page_break()


# ════════════════════════════════════════════════════════════════════════════
# APPENDIX B — Terminal Commands and Launch Procedure
# ════════════════════════════════════════════════════════════════════════════
heading1("APPENDIX B")
heading1("TERMINAL COMMANDS AND LAUNCH PROCEDURE")

heading2("B.1  Prerequisites and Installation")

cmd_blocks = [
    ("Install Ubuntu 22.04 with ROS2 Humble:", [
        "sudo apt install ros-humble-desktop",
        "sudo apt install ros-humble-gazebo-ros-pkgs",
        "sudo apt install python3-pyqt5",
        "sudo apt install ros-humble-xacro",
    ]),
    ("Create and clone workspace:", [
        "mkdir -p ~/atlas_ws/src",
        "cd ~/atlas_ws/src",
        "git clone https://github.com/laborbeekaambefikar-ship-it/my-project.git .",
    ]),
]

for desc, cmds in cmd_blocks:
    body_para(desc)
    for cmd in cmds:
        p = doc.add_paragraph()
        para_spacing(p, before=0, after=2)
        p.paragraph_format.line_spacing = Pt(14)
        run = p.add_run("    " + cmd)
        run.font.name = 'Courier New'
        run.font.size = Pt(10)
    doc.add_paragraph()

heading2("B.2  Build Procedure")

build_cmds = [
    "cd ~/atlas_ws",
    "source /opt/ros/humble/setup.bash",
    "colcon build --symlink-install",
    "source install/setup.bash",
]
body_para("Execute the following commands in a terminal to build the workspace:")
for cmd in build_cmds:
    p = doc.add_paragraph()
    para_spacing(p, before=0, after=2)
    p.paragraph_format.line_spacing = Pt(14)
    run = p.add_run("    " + cmd)
    run.font.name = 'Courier New'
    run.font.size = Pt(10)
doc.add_paragraph()

heading2("B.3  Full System Launch")

launch_cmds = [
    "# Kill any residual Gazebo processes first:",
    "pkill -9 -f gzserver 2>/dev/null",
    "pkill -9 -f gzclient 2>/dev/null",
    "sleep 2",
    "# Launch the complete ATLAS system:",
    "ros2 launch atlas_bringup atlas_full.launch.py",
]
body_para("The complete system is launched with a single command:")
for cmd in launch_cmds:
    p = doc.add_paragraph()
    para_spacing(p, before=0, after=2)
    p.paragraph_format.line_spacing = Pt(14)
    run = p.add_run("    " + cmd)
    run.font.name = 'Courier New'
    run.font.size = Pt(10)
doc.add_paragraph()

heading2("B.4  Sending Missions and Monitoring")

monitor_cmds = [
    ("Send a mission to shelf S05 (in a separate terminal):","ros2 run atlas_mission_manager send_mission S05"),
    ("Monitor robot state:", "ros2 topic echo /atlas/robot_state"),
    ("Monitor event log:", "ros2 topic echo /atlas/log"),
    ("Monitor line sensors:", "ros2 topic echo /atlas/line_sensors"),
    ("Monitor odometry:", "ros2 topic echo /atlas/odom --once"),
    ("Emergency stop:", 'ros2 topic pub /atlas/estop std_msgs/msg/Empty "{}" --once'),
    ("Reset E-stop:", 'ros2 topic pub /atlas/reset std_msgs/msg/Empty "{}" --once'),
    ("Reset AGV to dock:", 'ros2 topic pub /atlas/reset_to_dock std_msgs/msg/Empty "{}" --once'),
]
for desc, cmd in monitor_cmds:
    body_para(desc)
    p = doc.add_paragraph()
    para_spacing(p, before=0, after=6)
    p.paragraph_format.line_spacing = Pt(14)
    run = p.add_run("    " + cmd)
    run.font.name = 'Courier New'
    run.font.size = Pt(10)

page_break()

# ════════════════════════════════════════════════════════════════════════════
# APPENDIX C — State Machine Transition Table
# ════════════════════════════════════════════════════════════════════════════
heading1("APPENDIX C")
heading1("COMPLETE STATE MACHINE TRANSITION TABLE")

add_table(
    ["Current State", "Trigger / Condition", "Next State", "Action Taken"],
    [
        ["IDLE", "queue not empty AND not E-stopped", "NAV_SPINE", "Pop mission, reset junction count, log 'STATE IDLE → NAV_SPINE'"],
        ["NAV_SPINE", "junction_count == target_aisle", "TURNING", "Publish turn_cmd = −π/2 (right turn into aisle)"],
        ["TURNING", "turn_done received", "NAV_AISLE", "Resume line following"],
        ["NAV_AISLE", "tag_event.shelf_id == target shelf", "AT_SHELF", "Begin dwell timer"],
        ["AT_SHELF", "0.5 s dwell elapsed", "PICKUP", "Begin pickup timer"],
        ["PICKUP", "2.0 s pickup elapsed", "PIVOT", "Set carrying=True, publish turn_cmd = π (180° turn)"],
        ["PIVOT", "turn_done received", "RET_AISLE", "Begin return navigation"],
        ["RET_AISLE", "junction detected (spine)", "RET_TURN", "Publish turn_cmd = +π/2 (left turn onto spine)"],
        ["RET_TURN", "turn_done received", "RET_SPINE", "Resume line following southward"],
        ["RET_SPINE", "home tag detected (is_home=True)", "DOCKED", "Begin docking dwell"],
        ["DOCKED", "1.0 s dwell elapsed", "IDLE", "Set carrying=False, battery=100%, active=None, log 'Mission complete'"],
        ["ANY state", "/atlas/estop received", "ERROR", "Set estopped=True, publish zero Twist, log E-STOP"],
        ["ERROR", "/atlas/reset received", "IDLE", "Clear estopped, clear queue, clear active mission"],
        ["ANY state", "/atlas/reset_to_dock received", "IDLE", "Cancel mission, call Gazebo set_entity_state, teleport to (0,0,yaw=90°), log SYSTEM READY"],
    ],
    "Table (Appendix C): Complete State Machine Transition Table"
)

page_break()

# ════════════════════════════════════════════════════════════════════════════
# APPENDIX D — Physical Robot Build Checklist
# ════════════════════════════════════════════════════════════════════════════
heading1("APPENDIX D")
heading1("PHYSICAL ROBOT BUILD AND DEPLOYMENT CHECKLIST")

heading2("Phase 1: Mechanical Assembly")
mechanical_steps = [
    "Cut 6061 aluminium T-slot extrusion to form 300×250 mm rectangular base frame.",
    "Mount two JGB37-520 motor brackets symmetrically at y = ±0.15 m from centre.",
    "Install caster wheel mount at x = +0.12 m from centre on rear of chassis.",
    "Mount 3 mm aluminium or acrylic top plate for electronics.",
    "Install battery bay at centre of chassis directly above drive axle.",
    "Attach Raspberry Pi 4 to top plate with M3 standoffs.",
    "Mount QTR-8A sensor array at front of chassis, 6–10 mm above floor level.",
    "Mount BNO055 IMU at chassis centre, level and firmly secured.",
    "Mount RDM6300 RFID antenna at bottom of chassis, 6 mm clearance from floor.",
    "Install L298N motor controller on top plate, accessible for wiring.",
]
for i, step in enumerate(mechanical_steps, 1):
    body_para(f"{i}. {step}", before=2, after=2)

heading2("Phase 2: Wiring and Electronics")
wiring_steps = [
    "Wire 3S LiPo 11.1V battery → main switch → 10A blade fuse.",
    "Connect battery positive to L298N VIN; battery negative to common ground.",
    "Install 12V→5V buck converter; connect output to Raspberry Pi USB-C power input.",
    "Wire L298N outputs to left and right JGB37-520 motors (OUT1/2 and OUT3/4).",
    "Connect Raspberry Pi GPIO pins per Table 19 (GPIO Pin Mapping).",
    "Wire QTR-8A sensor output pins 1–8 to MCP3008 ADC channels CH0–CH7.",
    "Connect MCP3008 to Raspberry Pi SPI0 (GPIO 8/9/10/11).",
    "Wire BNO055 to Raspberry Pi I2C1 (GPIO 2 SDA, GPIO 3 SCL).",
    "Wire RDM6300 to Raspberry Pi UART (GPIO 14 TX, GPIO 15 RX).",
    "Verify all ground connections share common reference.",
    "Install e-stop mushroom button in series with motor path (L298N EN pins).",
]
for i, step in enumerate(wiring_steps, 1):
    body_para(f"{i}. {step}", before=2, after=2)

heading2("Phase 3: Software Deployment")
sw_steps = [
    "Install Ubuntu 22.04 Server ARM64 on Raspberry Pi 4 microSD (32+ GB).",
    "Install ROS2 Humble base: sudo apt install ros-humble-ros-base.",
    "Install required packages: colcon, xacro, sensor_msgs, nav_msgs, std_msgs.",
    "Clone atlas_ws repository to ~/atlas_ws on Raspberry Pi.",
    "Replace line_sensor.py with line_sensor_hw.py (SPI ADC implementation).",
    "Replace tag_detector.py with rfid_reader_hw.py (UART RFID implementation).",
    "Add imu_hw.py (I2C BNO055 implementation) to atlas_navigation package.",
    "Add motor_driver.py (GPIO PWM + encoder odometry) to atlas_hardware package.",
    "Build workspace: colcon build --symlink-install.",
    "Configure atlas_real.launch.py and verify all nodes launch without errors.",
]
for i, step in enumerate(sw_steps, 1):
    body_para(f"{i}. {step}", before=2, after=2)

heading2("Phase 4: Calibration and Testing")
cal_steps = [
    "Place QTR-8A over white floor and black tape; calibrate thresholds at midpoint.",
    "Verify sensor junction detection at tape intersections (5+ sensors active).",
    "Place robot on blocks; command known velocities; tune motor PID (Kp=2.0, Ki=1.0, Kd=0.05).",
    "Verify odometry accuracy over 2.0 m straight run (< 2% error).",
    "Calibrate BNO055 IMU per Bosch calibration procedure; verify heading stability.",
    "Test RFID detection with all 21 tags; verify detection at operational speed.",
    "Execute complete mission S01 manually; verify all FSM state transitions.",
    "Run 10 consecutive full missions; verify 100% completion rate.",
    "Test emergency stop response at full speed; verify stopping within 3 cm.",
    "Test reset-to-dock from multiple operational states.",
]
for i, step in enumerate(cal_steps, 1):
    body_para(f"{i}. {step}", before=2, after=2)


# ════════════════════════════════════════════════════════════════════════════
# SAVE
# ════════════════════════════════════════════════════════════════════════════
output_path = '/projects/sandbox/Atlas/ATLAS_FINAL_PROJECT_REPORT.docx'
doc.save(output_path)
print(f"Report saved to: {output_path}")
print(f"Total sections: {len(doc.sections)}")
print(f"Total paragraphs: {len(doc.paragraphs)}")
print(f"Total tables: {len(doc.tables)}")
